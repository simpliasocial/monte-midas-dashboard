from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches


REQUESTED_DOCX_PATH = Path(r"C:\Users\cristian.gomezcoello\Downloads\Implementacion_Supabase_SimpliaLeads_ISO10013_v1_2.docx")
REQUESTED_SQL_PATH = Path(r"C:\Users\cristian.gomezcoello\Downloads\supabase_replicacion_simpliale_solo_public_cw_actualizado_20260528.sql")

DOCX_CANDIDATES = [
    REQUESTED_DOCX_PATH,
    Path(r"D:\TRABAJO SIMPLIA\repositorios simplia.social\testings-dashboard\docs\Implementacion_Supabase_SimpliaLeads_ISO10013_v1_2.docx"),
]
SQL_CANDIDATES = [
    REQUESTED_SQL_PATH,
    Path(r"D:\TRABAJO SIMPLIA\repositorios simplia.social\testings-dashboard\docs\supabase_replicacion_simpliale_solo_public_cw.sql"),
]


def resolve_existing_path(candidates: list[Path], label: str) -> Path:
    for candidate in candidates:
        if candidate.exists():
            return candidate
    raise FileNotFoundError(f"No se encontro {label}. Rutas revisadas: {', '.join(str(path) for path in candidates)}")


DOCX_PATH = resolve_existing_path(DOCX_CANDIDATES, "documento Word de implementacion")
SQL_PATH = resolve_existing_path(SQL_CANDIDATES, "SQL de replicacion")

DOCX_MARKER = "Actualización v2.2 - Implementacion verificada al 04/06/2026"
DOCX_META_ADS_MARKER = "Actualización v2.3 - Meta Ads configurable por negocio al 04/06/2026"
SQL_MARKER = "INICIO BLOQUE 12: addendum 2026-06-04 labels vivos, pruning y discovery"
SQL_META_ADS_MARKER = "INICIO BLOQUE 16: Meta Ads configurable por negocio"


def backup(path: Path) -> Path:
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_path = path.with_name(f"{path.stem}.bak_{stamp}{path.suffix}")
    backup_path.write_bytes(path.read_bytes())
    return backup_path


def add_paragraphs(document: Document, items: list[str], style: str | None = None) -> None:
    for item in items:
        document.add_paragraph(item, style=style)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    document.add_paragraph()


def update_docx() -> bool:
    document = Document(DOCX_PATH)
    full_text = "\n".join(p.text for p in document.paragraphs)
    if DOCX_MARKER in full_text:
        return False

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(DOCX_MARKER, level=1)

    document.add_heading("23.1 Objetivo Del Addendum", level=2)
    add_paragraphs(document, [
        "Este addendum documenta el estado real implementado en Supabase y el dashboard despues de la puesta en marcha con Chatwoot, importacion historica de n8n, limpieza de conversaciones legacy y sincronizacion diaria. Su proposito es que el mismo modelo pueda replicarse en otra compania cambiando solamente variables, credenciales y catalogos propios del negocio.",
        "La base tecnica que debe permanecer estable entre negocios es: schema cw, tablas de snapshot e historico, Edge Functions, Vault, pg_cron, webhooks de Chatwoot, discovery automatico de etiquetas y atributos, importacion historica opcional y pruebas de salud.",
        "No se deben copiar claves reales entre negocios. Toda API key, token, remitente de correo, cuenta Chatwoot, proyecto Supabase, proveedor de WhatsApp/Zapaway/Zapi u otro proveedor externo debe configurarse como placeholder y reemplazarse por valores verificados del nuevo negocio.",
    ])

    document.add_heading("23.2 Cambios Nuevos Integrados", level=2)
    add_table(document, ["Cambio", "Objeto", "Para que sirve al replicar"], [
        ["Aislamiento por cuenta Chatwoot", "chatwoot_account_id en conversaciones/mensajes y validacion del dashboard", "Evita que un ID de conversacion antiguo se convierta en link de la cuenta actual. Si el registro no pertenece al Chatwoot del negocio, no se abre como conversacion viva."],
        ["Labels vivos desde Chatwoot", "cw.label_catalog + refresh_dashboard_discovery", "El dashboard toma etiquetas reales del Chatwoot actual. Al borrar una etiqueta en Chatwoot, no debe seguir apareciendo como opcion operativa."],
        ["Pruning de etiquetas eliminadas", "cw.prune_dashboard_settings_labels y cw.prune_deleted_label_references", "Limpia settings, conversaciones y eventos para que no queden etiquetas fantasma."],
        ["Webhook de labels/atributos", "chatwoot-label-webhook", "Recibe cambios de Chatwoot, registra deltas de etiquetas, actualiza snapshot de conversacion y refresca discovery."],
        ["Sync diario verificado", "chatwoot-sync + cron sync-chatwoot-diario", "Ejecuta sincronizacion diaria a las 00:01 Ecuador usando pg_cron y pg_net."],
        ["Historico n8n ordenado", "source_system = n8n_chat_histories", "Permite conservar data historica importada sin hacerla pasar por conversaciones vivas de Chatwoot."],
        ["Archivo de legacy Chatwoot", "cw.archived_legacy_chatwoot_*", "Antes de borrar conversaciones de una cuenta Chatwoot equivocada, se archivan filas y mensajes para auditoria."],
        ["Secrets por ambiente", "Supabase Edge Function Secrets + Vault", "Mantiene secretos server-side. El frontend solo conserva URL y anon key publicas cuando aplica."],
    ])

    document.add_heading("23.3 Variables Que Cambian Por Negocio", level=2)
    add_table(document, ["Variable", "Donde se configura", "Nota"], [
        ["VITE_SUPABASE_URL", ".env frontend y Edge Function Secret SUPABASE_URL", "URL del proyecto Supabase del negocio."],
        ["VITE_SUPABASE_ANON_KEY", ".env frontend", "Clave publica anon/publishable. No usar service role en frontend."],
        ["SUPABASE_SERVICE_ROLE_KEY", "Edge Function Secret", "Solo server-side para funciones que escriben en cw."],
        ["VITE_CHATWOOT_BASE_URL / CHATWOOT_BASE_URL", ".env y Edge Function Secrets", "Normalmente https://app.chatwoot.com o instancia self-hosted."],
        ["VITE_CHATWOOT_ACCOUNT_ID / CHATWOOT_ACCOUNT_ID", ".env y Edge Function Secrets", "Cuenta Chatwoot del negocio. Es la llave para evitar mezclar companias."],
        ["VITE_CHATWOOT_API_TOKEN / CHATWOOT_API_TOKEN", ".env si aplica y Edge Function Secrets", "Token API de Chatwoot del usuario/bot de automatizacion."],
        ["CHATWOOT_WEBHOOK_SECRET", "Edge Function Secret y URL/header del webhook", "Se usa para autorizar chatwoot-label-webhook."],
        ["OPENAI_API_KEY", "Edge Function Secret", "Nunca en frontend. Usado por reportes IA."],
        ["OPENAI_REPORT_MODEL", "Edge Function Secret opcional", "Modelo de reportes. Mantener parametrizable por costo/calidad."],
        ["RESEND_API_KEY", "Edge Function Secret", "Necesaria para correos programados."],
        ["RESEND_FROM_EMAIL", "Edge Function Secret", "Debe ser remitente verificado, por ejemplo Nombre <correo@dominio-verificado.com>."],
        ["ZAPAWAY_API_KEY / proveedor WhatsApp", "Secret del proveedor si el negocio lo usa", "No forma parte del schema cw base, pero debe documentarse como integracion externa por compania."],
        ["DASHBOARD_ACCOUNT_ID", "Body del cron o Edge Function Secret opcional", "Por defecto 0. Usar otro valor si el negocio maneja multiples cuentas logicas dentro del mismo proyecto."],
    ])

    document.add_heading("23.4 Flujo Replicable De Implementacion", level=2)
    add_paragraphs(document, [
        "1. Crear proyecto Supabase del nuevo negocio y confirmar project_ref, region y acceso CLI.",
        "2. Configurar .env del frontend con URL/anon key del nuevo proyecto y variables Chatwoot del nuevo negocio.",
        "3. Ejecutar SQL base por bloques 01-11.",
        "4. Ejecutar el addendum SQL 2026-06-04: funciones de pruning, validaciones de labels vivos y bloque opcional de archivo legacy.",
        "5. Guardar Edge Function Secrets. Mantener alias VITE_CHATWOOT_* y CHATWOOT_* si las funciones actuales los consumen.",
        "6. Guardar Vault secrets para pg_cron: project_url y JWT/token server-side que pueda invocar Edge Functions.",
        "7. Desplegar Edge Functions: chatwoot-sync, chatwoot-label-webhook, chatwoot-repair-conversations, generate-ai-report, send-scheduled-reports y las funciones extra que aplique el negocio.",
        "8. Configurar webhook en Chatwoot apuntando a /functions/v1/chatwoot-label-webhook?secret=<CHATWOOT_WEBHOOK_SECRET>. Activar eventos de actualizacion de conversacion, labels y contacto cuando esten disponibles.",
        "9. Ejecutar un sync full manual con ventana inicial. Validar cw.sync_runs en success y conteos en cw.conversations_current, cw.messages, cw.label_catalog y cw.attribute_definitions.",
        "10. Si existe historico n8n, ejecutar bloque 10. El historico debe quedar con source_system = n8n_chat_histories y no debe generar links Chatwoot vivos si no pertenece a la cuenta actual.",
        "11. Si el repositorio venia de otra compania, ejecutar el bloque opcional de archivo legacy con el CHATWOOT_ACCOUNT_ID actual. Confirmar que active_legacy_conversations = 0.",
        "12. Correr npm run check y una prueba funcional del dashboard: pestaña Conversaciones, filtros de etiquetas, exportes y boton Abrir conversacion.",
    ])

    document.add_heading("23.5 Chatwoot Sync Y Cron", level=2)
    add_paragraphs(document, [
        "La Edge Function chatwoot-sync sincroniza inboxes, labels, attribute_definitions, contacts_current, conversations_current, messages, sync_cursor y sync_runs. En mensajes, si Chatwoot no envia account_id dentro del mensaje, la funcion debe heredar conv.account_id para mantener trazabilidad por cuenta.",
        "El cron operativo se llama sync-chatwoot-diario y usa schedule 1 5 * * *, equivalente a 00:01 America/Guayaquil porque pg_cron opera en UTC. El body recomendado es {\"source\":\"pg_cron\",\"mode\":\"full\",\"window_hours\":72,\"sync_messages\":\"recent\",\"dashboard_account_id\":0}.",
        "La evidencia de ejecucion se revisa en cron.job_run_details y cw.sync_runs. Una corrida correcta debe terminar con status success y error_message vacio.",
    ])

    document.add_heading("23.6 Labels, Discovery Y Pruning", level=2)
    add_paragraphs(document, [
        "El dashboard ya no debe depender de una lista fija de etiquetas. La fuente viva es cw.label_catalog, alimentada por chatwoot-sync y chatwoot-label-webhook.",
        "cw.refresh_dashboard_discovery actualiza availableLabels, discoveredLabels, availableAttributeKeys y autoTagGroups. cw.prune_dashboard_settings_labels elimina de la configuracion cualquier etiqueta que ya no exista en Chatwoot. cw.prune_deleted_label_references limpia labels en conversations_current y conversation_label_events.",
        "Cada vez que se borre o renombre una etiqueta en Chatwoot, ejecutar sync o recibir webhook debe dejar el dashboard sin etiquetas fantasma.",
    ])

    document.add_heading("23.7 Limpieza De Conversaciones Legacy", level=2)
    add_paragraphs(document, [
        "Caso detectado: existian conversaciones antiguas de otro Chatwoot con IDs 1-7. Al cambiar el account_id en .env, el dashboard armaba links del nuevo Chatwoot usando IDs viejos. Esto se corrige en dos capas: base de datos y frontend.",
        "Base de datos: archivar y borrar de tablas activas cualquier conversacion cuyo chatwoot_account_id o raw_payload.account_id no coincida con el CHATWOOT_ACCOUNT_ID actual. El archivo queda en cw.archived_legacy_chatwoot_* con RLS activo y sin grants a anon/authenticated.",
        "Frontend: getChatwootUrl debe recibir el objeto de conversacion completo. Si chatwoot_account_id no coincide con config.chatwoot.accountId, retorna vacio y el boton se muestra deshabilitado como Sin conversacion en Chatwoot.",
    ])

    document.add_heading("23.8 Validaciones Obligatorias Para Otro Negocio", level=2)
    add_table(document, ["Validacion", "Resultado esperado"], [
        ["functions list", "chatwoot-sync y chatwoot-label-webhook ACTIVE."],
        ["cron.job", "sync-chatwoot-diario active = true, schedule = 1 5 * * *, invoca /functions/v1/chatwoot-sync."],
        ["cw.sync_runs", "Ultima corrida status = success, error_message vacio."],
        ["cw.sync_cursor", "daily_full actualizado con ventana reciente."],
        ["cw.label_catalog", "Cantidad de labels igual al catalogo real del Chatwoot del negocio."],
        ["stale labels", "0 etiquetas fuera de cw.label_catalog en conversations_current y conversation_label_events."],
        ["legacy Chatwoot", "0 conversaciones activas con account_id distinto al actual."],
        ["historico n8n", "Registros preservados con source_system = n8n_chat_histories."],
        ["dashboard", "Abrir conversacion solo habilitado para registros que pertenecen al Chatwoot actual."],
    ])

    document.add_heading("23.9 Comandos Operativos Con Placeholders", level=2)
    add_paragraphs(document, [
        "Login y link CLI: npx supabase login --token \"<SUPABASE_ACCESS_TOKEN>\"; npx supabase link --project-ref <PROJECT_REF>",
        "Deploy functions: npx supabase functions deploy chatwoot-sync chatwoot-repair-conversations generate-ai-report meta-campaign-insights send-scheduled-reports --project-ref <PROJECT_REF> --use-api --jobs 2",
        "Deploy webhook sin JWT: npx supabase functions deploy chatwoot-label-webhook --project-ref <PROJECT_REF> --use-api --no-verify-jwt",
        "Sync manual por HTTP: POST https://<PROJECT_REF>.supabase.co/functions/v1/chatwoot-sync con Authorization Bearer <ANON_OR_SERVICE_JWT> y body {\"mode\":\"full\",\"window_hours\":72,\"sync_messages\":\"recent\",\"dashboard_account_id\":0}.",
        "Webhook Chatwoot: https://<PROJECT_REF>.supabase.co/functions/v1/chatwoot-label-webhook?secret=<CHATWOOT_WEBHOOK_SECRET>.",
    ])

    document.add_heading("23.10 Criterio De Cierre", level=2)
    add_paragraphs(document, [
        "La replica se considera lista cuando: SQL aplicado sin errores, secrets configurados, Edge Functions ACTIVE, cron activo, sync_runs success, labels vivos sin fantasma, historico importado preservado, legacy externo en cero, reportes IA/envio de correo probados y npm run check pasa en el frontend.",
    ])

    document.save(DOCX_PATH)
    return True


def update_docx_meta_ads() -> bool:
    document = Document(DOCX_PATH)
    full_text = "\n".join(p.text for p in document.paragraphs)
    if DOCX_META_ADS_MARKER in full_text:
        return False

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(DOCX_META_ADS_MARKER, level=1)

    document.add_heading("24.1 Objetivo Del Addendum Meta Ads", level=2)
    add_paragraphs(document, [
        "Este addendum documenta la integracion nueva para configurar Meta Ads desde el dashboard por cada negocio. El objetivo es que la pestana Tendencias pueda traer campanas e insights sin hardcodear el ID de cuenta publicitaria ni exponer el Bearer Token en el frontend.",
        "El flujo replicable queda asi: Platform Admin o Company Admin abre Tendencias > Meta Ads > Configurar campanas, ingresa el Ad Account ID y el Bearer Token, Supabase guarda la configuracion server-side y la Edge Function meta-campaign-insights usa esos valores para consultar Graph API.",
        "Operators no ven el boton de configuracion. Pueden seguir consumiendo las vistas autorizadas del dashboard segun los permisos definidos por rol.",
    ])

    document.add_heading("24.2 Componentes Implementados", level=2)
    add_table(document, ["Componente", "Archivo u objeto", "Responsabilidad"], [
        ["Tabla segura de configuracion", "cw.meta_ads_configs", "Guarda account_id, ad_account_id, access_token, token_last_four, graph_api_version, enabled y auditoria de configuracion."],
        ["RLS y permisos", "cw.meta_ads_configs", "RLS activo. Sin permisos para anon/authenticated. Solo service_role/postgres puede leer o escribir el token completo."],
        ["Edge Function", "meta-campaign-insights", "Expone acciones get_config, save_config y fetch_insights. Autentica usuario, valida rol y consulta Meta con token server-side."],
        ["Cliente frontend", "MetaAdsInsightsClient", "Invoca la funcion. El frontend solo recibe configuracion publica: configured, adAccountId, tokenLast4 y version."],
        ["Hook React", "useMetaAdsCampaignConfig", "Carga configuracion, guarda cambios e invalida cache de insights."],
        ["Dialog UI", "MetaCampaignConfigDialog", "Formulario para Ad Account ID, Bearer Token y Graph API version. El token no se persiste en localStorage."],
        ["Permiso por rol", "canConfigureMetaAds", "Permite configurar a platform_admin y company_admin. Rechaza operator."],
    ])

    document.add_heading("24.3 Seguridad Del Bearer Token", level=2)
    add_paragraphs(document, [
        "Regla critica: el Bearer Token de Meta Ads nunca debe configurarse como variable VITE_* ni guardarse en localStorage. Tampoco debe devolverse completo al navegador.",
        "La tabla cw.meta_ads_configs almacena access_token server-side. La UI solo muestra token_last_four para confirmar rotacion o existencia del token.",
        "La Edge Function usa SUPABASE_SERVICE_ROLE_KEY para leer la configuracion. Antes de guardar o leer, valida el JWT del usuario y consulta public.user_profiles.role. No usa user_metadata como fuente de autorizacion.",
        "Existe fallback legacy: si no hay fila en cw.meta_ads_configs, meta-campaign-insights puede usar META_AD_ACCOUNT_ID, META_SYSTEM_USER_TOKEN y META_GRAPH_API_VERSION desde Supabase Secrets. Esto permite migrar proyectos antiguos sin romper la carga inicial.",
    ])

    document.add_heading("24.4 Variables Que Cambian Por Negocio", level=2)
    add_table(document, ["Variable", "Donde se configura", "Observacion"], [
        ["Meta Ad Account ID", "Dashboard > Tendencias > Meta Ads > Configurar campanas", "Puede pegarse con o sin prefijo act_. La funcion lo normaliza a valor numerico."],
        ["Meta Bearer Token", "Dashboard > Tendencias > Meta Ads > Configurar campanas", "Debe tener permisos de lectura de campañas e insights para la cuenta publicitaria."],
        ["Graph API Version", "Dashboard o fallback META_GRAPH_API_VERSION", "Por defecto v20.0. Mantener parametrizable por negocio."],
        ["META_AD_ACCOUNT_ID", "Supabase Secret opcional legacy", "Solo fallback. La ruta recomendada es cw.meta_ads_configs."],
        ["META_SYSTEM_USER_TOKEN", "Supabase Secret opcional legacy", "Solo fallback. Nunca frontend."],
        ["META_CACHE_TTL_SECONDS", "Supabase Secret opcional", "Controla cache de insights. Valor usado: minimo 60 segundos; sugerido 900."],
    ])

    document.add_heading("24.5 Flujo De Operacion", level=2)
    add_paragraphs(document, [
        "1. Ejecutar el Bloque 16 del SQL de replicacion para crear cw.meta_ads_configs con RLS y permisos seguros.",
        "2. Desplegar la Edge Function actualizada: npx supabase functions deploy meta-campaign-insights --project-ref <PROJECT_REF> --use-api.",
        "3. Iniciar sesion en el dashboard con un usuario platform_admin o company_admin.",
        "4. Ir a Tendencias > Meta Ads y abrir Configurar campanas.",
        "5. Guardar Ad Account ID y Bearer Token. Si ya existe token guardado, dejar el campo vacio conserva el token anterior.",
        "6. Pulsar Actualizar. La funcion usa el token server-side, consulta campaigns e insights y escribe cache en cw.meta_campaigns_current, cw.meta_adset_insights_cache y cw.meta_ads_sync_runs.",
        "7. Revisar que el badge de la tarjeta muestre la cuenta configurada y que cw.meta_ads_sync_runs registre success o cache_hit segun corresponda.",
    ])

    document.add_heading("24.6 Validaciones Recomendadas", level=2)
    add_table(document, ["Validacion", "Resultado esperado"], [
        ["RLS meta_ads_configs", "relrowsecurity = true."],
        ["Grants meta_ads_configs", "No existen grants para anon ni authenticated; service_role tiene select/insert/update/delete."],
        ["Funcion desplegada", "meta-campaign-insights aparece ACTIVE en Supabase Functions."],
        ["Permisos UI", "platform_admin y company_admin ven Configurar campanas; operator no lo ve."],
        ["Config publica", "get_config devuelve configured, adAccountId, tokenLast4 y graphApiVersion; nunca access_token."],
        ["Cache Meta Ads", "Despues de Actualizar existen filas en cw.meta_campaigns_current y/o cw.meta_adset_insights_cache segun entrega de Meta."],
        ["Sync runs", "cw.meta_ads_sync_runs registra status success/cache_hit/error con fechas del rango solicitado."],
    ])

    document.add_heading("24.7 SQL Y Runbook Replicable", level=2)
    add_paragraphs(document, [
        "El SQL de replicacion incluye ahora el Bloque 16: Meta Ads configurable por negocio. Este bloque es idempotente y se puede ejecutar despues del esquema base de cw.",
        "No se deben insertar tokens reales en el SQL. La configuracion se realiza desde el dashboard o, en escenarios automatizados, invocando la Edge Function save_config con un usuario administrador autenticado.",
        "Para replicar en otra empresa, solo cambian credenciales y cuenta Meta. La estructura cw.meta_ads_configs, meta_campaigns_current, meta_adset_insights_cache y meta_ads_sync_runs se mantiene igual.",
    ])

    document.save(DOCX_PATH)
    return True


SQL_ADDENDUM = r"""

-- ========================================================================
-- INICIO BLOQUE 12: addendum 2026-06-04 labels vivos, pruning y discovery
-- DONDE SE EJECUTA: Supabase > SQL Editor
-- INSTRUCCION: ejecutar despues de los bloques 01-11. Es idempotente.
-- OBJETIVO: que el dashboard use etiquetas reales de Chatwoot, limpie
-- etiquetas eliminadas y refresque discovery para otro negocio.
-- ========================================================================

create or replace function cw.prune_dashboard_settings_labels(target_account_id bigint default 0)
returns jsonb
language plpgsql
security definer
set search_path = cw, pg_catalog
as $$
declare
    current_settings jsonb;
    pruned_settings jsonb;
    label_list text[];
    key_name text;
    label_array_keys text[] := array[
        'sqlTags',
        'appointmentTags',
        'saleTags',
        'unqualifiedTags',
        'scoreHighTags',
        'scoreMediumTags',
        'scoreLowTags',
        'humanFollowupQueueTags',
        'humanSalesQueueTags',
        'scoreAppointmentLabels'
    ];
    auto_label_array_keys text[] := array[
        'sqlTags',
        'appointmentTags',
        'saleTags',
        'unqualifiedTags',
        'humanFollowupQueueTags'
    ];
begin
    select settings
    into current_settings
    from cw.dashboard_tag_settings
    where account_id = target_account_id;

    select coalesce(array_agg(title order by title), array[]::text[])
    into label_list
    from cw.label_catalog
    where account_id = target_account_id;

    pruned_settings := coalesce(current_settings, '{}'::jsonb)
        || jsonb_build_object(
            'availableLabels', to_jsonb(label_list),
            'discoveredLabels', to_jsonb(label_list)
        );

    foreach key_name in array label_array_keys loop
        pruned_settings := jsonb_set(
            pruned_settings,
            array[key_name],
            (
                select coalesce(jsonb_agg(item_label order by item_label), '[]'::jsonb)
                from (
                    select distinct item_label
                    from jsonb_array_elements_text(
                        case
                            when jsonb_typeof(pruned_settings -> key_name) = 'array'
                                then pruned_settings -> key_name
                            else '[]'::jsonb
                        end
                    ) as labels(item_label)
                    where item_label = any(label_list)
                ) kept
            ),
            true
        );
    end loop;

    if not (coalesce(pruned_settings ->> 'humanAppointmentTargetLabel', '') = any(label_list)) then
        pruned_settings := jsonb_set(pruned_settings, '{humanAppointmentTargetLabel}', to_jsonb(''::text), true);
    end if;

    if not (coalesce(pruned_settings ->> 'humanSaleTargetLabel', '') = any(label_list)) then
        pruned_settings := jsonb_set(pruned_settings, '{humanSaleTargetLabel}', to_jsonb(''::text), true);
    end if;

    if jsonb_typeof(pruned_settings -> 'autoTagGroups') = 'object' then
        foreach key_name in array auto_label_array_keys loop
            pruned_settings := jsonb_set(
                pruned_settings,
                array['autoTagGroups', key_name],
                (
                    select coalesce(jsonb_agg(item_label order by item_label), '[]'::jsonb)
                    from (
                        select distinct item_label
                        from jsonb_array_elements_text(
                            case
                                when jsonb_typeof(pruned_settings #> array['autoTagGroups', key_name]) = 'array'
                                    then pruned_settings #> array['autoTagGroups', key_name]
                                else '[]'::jsonb
                            end
                        ) as labels(item_label)
                        where item_label = any(label_list)
                    ) kept
                ),
                true
            );
        end loop;
    end if;

    insert into cw.dashboard_tag_settings (account_id, settings, updated_at)
    values (target_account_id, pruned_settings, now())
    on conflict (account_id)
    do update set settings = excluded.settings, updated_at = now();

    return jsonb_build_object(
        'account_id', target_account_id,
        'labels', label_list,
        'settings', pruned_settings,
        'updated_at', now()
    );
end;
$$;

revoke all on function cw.prune_dashboard_settings_labels(bigint) from anon, authenticated;
grant execute on function cw.prune_dashboard_settings_labels(bigint) to service_role;

create or replace function cw.prune_deleted_label_references(target_account_id bigint default 0)
returns jsonb
language plpgsql
security definer
set search_path = cw, pg_catalog
as $$
declare
    active_labels text[];
    conversations_updated integer := 0;
    events_updated integer := 0;
    events_deleted integer := 0;
begin
    select coalesce(array_agg(title order by title), array[]::text[])
    into active_labels
    from cw.label_catalog
    where account_id = target_account_id;

    with pruned as (
        select
            c.chatwoot_conversation_id,
            coalesce((
                select array_agg(label order by label)
                from (
                    select distinct label
                    from unnest(coalesce(c.labels, array[]::text[])) as labels(label)
                    where label = any(active_labels)
                ) kept
            ), array[]::text[]) as labels
        from cw.conversations_current c
        where exists (
            select 1
            from unnest(coalesce(c.labels, array[]::text[])) as labels(label)
            where not (label = any(active_labels))
        )
    )
    update cw.conversations_current c
    set labels = pruned.labels,
        updated_at = now()
    from pruned
    where c.chatwoot_conversation_id = pruned.chatwoot_conversation_id;

    get diagnostics conversations_updated = row_count;

    with pruned as (
        select
            e.id,
            coalesce((
                select array_agg(label order by label)
                from (
                    select distinct label
                    from unnest(coalesce(e.previous_labels, array[]::text[])) as labels(label)
                    where label = any(active_labels)
                ) kept
            ), array[]::text[]) as previous_labels,
            coalesce((
                select array_agg(label order by label)
                from (
                    select distinct label
                    from unnest(coalesce(e.next_labels, array[]::text[])) as labels(label)
                    where label = any(active_labels)
                ) kept
            ), array[]::text[]) as next_labels,
            coalesce((
                select array_agg(label order by label)
                from (
                    select distinct label
                    from unnest(coalesce(e.added_labels, array[]::text[])) as labels(label)
                    where label = any(active_labels)
                ) kept
            ), array[]::text[]) as added_labels,
            coalesce((
                select array_agg(label order by label)
                from (
                    select distinct label
                    from unnest(coalesce(e.removed_labels, array[]::text[])) as labels(label)
                    where label = any(active_labels)
                ) kept
            ), array[]::text[]) as removed_labels
        from cw.conversation_label_events e
        where exists (
            select 1
            from unnest(
                coalesce(e.previous_labels, array[]::text[]) ||
                coalesce(e.next_labels, array[]::text[]) ||
                coalesce(e.added_labels, array[]::text[]) ||
                coalesce(e.removed_labels, array[]::text[])
            ) as labels(label)
            where not (label = any(active_labels))
        )
    )
    update cw.conversation_label_events e
    set previous_labels = pruned.previous_labels,
        next_labels = pruned.next_labels,
        added_labels = pruned.added_labels,
        removed_labels = pruned.removed_labels
    from pruned
    where e.id = pruned.id;

    get diagnostics events_updated = row_count;

    delete from cw.conversation_label_events
    where coalesce(array_length(previous_labels, 1), 0) = 0
      and coalesce(array_length(next_labels, 1), 0) = 0
      and coalesce(array_length(added_labels, 1), 0) = 0
      and coalesce(array_length(removed_labels, 1), 0) = 0;

    get diagnostics events_deleted = row_count;

    return jsonb_build_object(
        'account_id', target_account_id,
        'active_labels', active_labels,
        'conversations_updated', conversations_updated,
        'events_updated', events_updated,
        'events_deleted', events_deleted
    );
end;
$$;

revoke all on function cw.prune_deleted_label_references(bigint) from anon, authenticated;
grant execute on function cw.prune_deleted_label_references(bigint) to service_role;

create or replace function cw.refresh_dashboard_discovery(target_account_id bigint default 0)
returns jsonb
language plpgsql
set search_path = cw, pg_catalog, public
as $$
declare
    current_settings jsonb := '{}'::jsonb;
    refreshed_settings jsonb := '{}'::jsonb;
    label_list text[] := '{}'::text[];
    all_attribute_keys text[] := '{}'::text[];
    contact_attribute_keys text[] := '{}'::text[];
    conversation_attribute_keys text[] := '{}'::text[];
    resolved_attribute_keys text[] := '{}'::text[];
    auto_sql_tags text[] := '{}'::text[];
    auto_appointment_tags text[] := '{}'::text[];
    auto_sale_tags text[] := '{}'::text[];
    auto_unqualified_tags text[] := '{}'::text[];
    auto_followup_tags text[] := '{}'::text[];
    auto_appointment_field_keys text[] := '{}'::text[];
    auto_sale_field_keys text[] := '{}'::text[];
    suggested_score_key text := '';
begin
    select coalesce(settings, '{}'::jsonb)
    into current_settings
    from cw.dashboard_tag_settings
    where account_id = target_account_id;

    current_settings := coalesce(current_settings, '{}'::jsonb);

    select coalesce(array_agg(label order by lower(label)), '{}'::text[])
    into label_list
    from (
        select distinct btrim(title) as label
        from cw.label_catalog
        where account_id = target_account_id
          and btrim(coalesce(title, '')) <> ''
    ) labels;

    with keys(attribute_scope, attribute_key, source_name) as (
        select attribute_scope, attribute_key, 'attribute_definitions'
        from cw.attribute_definitions
        where btrim(coalesce(attribute_key, '')) <> ''
        union all
        select 'contact', jsonb_object_keys(custom_attributes), 'contacts_current.custom_attributes'
        from cw.contacts_current
        where jsonb_typeof(custom_attributes) = 'object'
        union all
        select 'contact', jsonb_object_keys(contact_custom_attributes), 'conversations_current.contact_custom_attributes'
        from cw.conversations_current
        where jsonb_typeof(contact_custom_attributes) = 'object'
        union all
        select 'conversation', jsonb_object_keys(conversation_custom_attributes), 'conversations_current.conversation_custom_attributes'
        from cw.conversations_current
        where jsonb_typeof(conversation_custom_attributes) = 'object'
        union all
        select 'resolved', jsonb_object_keys(custom_attributes), 'conversations_current.custom_attributes'
        from cw.conversations_current
        where jsonb_typeof(custom_attributes) = 'object'
    ),
    cleaned as (
        select
            case
                when attribute_scope in ('contact', 'conversation', 'resolved') then attribute_scope
                else 'unknown'
            end as attribute_scope,
            btrim(attribute_key) as attribute_key,
            source_name
        from keys
        where btrim(coalesce(attribute_key, '')) <> ''
    ),
    grouped as (
        select
            attribute_scope,
            attribute_key,
            array_agg(distinct source_name order by source_name) as source_names,
            count(*)::bigint as rows_seen
        from cleaned
        group by attribute_scope, attribute_key
    )
    insert into cw.attribute_key_catalog (
        account_id,
        attribute_scope,
        attribute_key,
        source_names,
        rows_seen,
        last_seen_at,
        updated_at
    )
    select
        target_account_id,
        attribute_scope,
        attribute_key,
        source_names,
        rows_seen,
        now(),
        now()
    from grouped
    on conflict (account_id, attribute_scope, normalized_key)
    do update set
        attribute_key = excluded.attribute_key,
        source_names = cw.merge_text_arrays(cw.attribute_key_catalog.source_names, excluded.source_names),
        rows_seen = excluded.rows_seen,
        last_seen_at = now(),
        updated_at = now();

    select coalesce(array_agg(attribute_key order by lower(attribute_key)), '{}'::text[])
    into contact_attribute_keys
    from (
        select distinct attribute_key
        from cw.attribute_key_catalog
        where account_id = target_account_id
          and attribute_scope = 'contact'
    ) keys;

    select coalesce(array_agg(attribute_key order by lower(attribute_key)), '{}'::text[])
    into conversation_attribute_keys
    from (
        select distinct attribute_key
        from cw.attribute_key_catalog
        where account_id = target_account_id
          and attribute_scope = 'conversation'
    ) keys;

    select coalesce(array_agg(attribute_key order by lower(attribute_key)), '{}'::text[])
    into resolved_attribute_keys
    from (
        select distinct attribute_key
        from cw.attribute_key_catalog
        where account_id = target_account_id
          and attribute_scope = 'resolved'
    ) keys;

    all_attribute_keys := cw.merge_text_arrays(
        cw.merge_text_arrays(contact_attribute_keys, conversation_attribute_keys),
        resolved_attribute_keys
    );

    with labels as (
        select label, cw.catalog_text_key(label) as key
        from unnest(label_list) as label
    )
    select
        coalesce(array_agg(label order by lower(label)) filter (
            where key ~ '(^|_)(interesado|interes|lead_calificado|calificado|qualified|sql|solicita|informacion|cotizacion|cotiza|prospecto|bienvenida)(_|$)'
        ), '{}'::text[]),
        coalesce(array_agg(label order by lower(label)) filter (
            where key ~ '(^|_)(cita|agendada|agendado|agenda|appointment|booking|reserva|reservado|visita|demo|reunion)(_|$)'
        ), '{}'::text[]),
        coalesce(array_agg(label order by lower(label)) filter (
            where key ~ '(^|_)(venta|vendido|vendida|sale|sold|won|ganada|ganado|cerrado|compra|comprado|pagado|facturado)(_|$)'
              and key !~ '(^|_)(no_venta|sin_venta|perdida|lost)(_|$)'
        ), '{}'::text[]),
        coalesce(array_agg(label order by lower(label)) filter (
            where key ~ '(^|_)(desinteresado|descartado|no_calificado|no_aplica|unqualified|rechazo|rechazado|perdido|perdida|lost|spam|invalido|invalid)(_|$)'
        ), '{}'::text[]),
        coalesce(array_agg(label order by lower(label)) filter (
            where key ~ '(^|_)(seguimiento|followup|follow_up|follow|pendiente)(_|$)'
        ), '{}'::text[])
    into
        auto_sql_tags,
        auto_appointment_tags,
        auto_sale_tags,
        auto_unqualified_tags,
        auto_followup_tags
    from labels;

    with attrs as (
        select attribute_key, cw.catalog_text_key(attribute_key) as key
        from unnest(all_attribute_keys) as attribute_key
    )
    select
        coalesce(array_agg(attribute_key order by lower(attribute_key)) filter (
            where key ~ '(^|_)(fecha_visita|hora_visita|fecha_cita|hora_cita|cita|visita|appointment|responsable|asesor|agente)(_|$)'
        ), '{}'::text[]),
        coalesce(array_agg(attribute_key order by lower(attribute_key)) filter (
            where key ~ '(^|_)(monto|monto_operacion|fecha_monto|venta|operacion|producto|servicio|responsable|asesor|agente)(_|$)'
        ), '{}'::text[])
    into auto_appointment_field_keys, auto_sale_field_keys
    from attrs;

    select attribute_key
    into suggested_score_key
    from unnest(all_attribute_keys) as attribute_key
    where cw.catalog_text_key(attribute_key) in ('score_interes', 'score', 'lead_score', 'puntaje')
    order by case cw.catalog_text_key(attribute_key)
        when 'score_interes' then 1
        when 'lead_score' then 2
        when 'score' then 3
        when 'puntaje' then 4
        else 5
    end
    limit 1;

    refreshed_settings := current_settings || jsonb_build_object(
        'autoDiscoveryEnabled', true,
        'lastAutoDiscoveryAt', to_jsonb(now()),
        'availableLabels', to_jsonb(label_list),
        'discoveredLabels', to_jsonb(label_list),
        'availableAttributeKeys', to_jsonb(all_attribute_keys),
        'discoveredAttributeKeys', to_jsonb(all_attribute_keys),
        'contactAttributeKeys', to_jsonb(contact_attribute_keys),
        'conversationAttributeKeys', to_jsonb(conversation_attribute_keys),
        'resolvedAttributeKeys', to_jsonb(resolved_attribute_keys),
        'autoTagGroups', jsonb_build_object(
            'sqlTags', to_jsonb(auto_sql_tags),
            'appointmentTags', to_jsonb(auto_appointment_tags),
            'saleTags', to_jsonb(auto_sale_tags),
            'unqualifiedTags', to_jsonb(auto_unqualified_tags),
            'humanFollowupQueueTags', to_jsonb(auto_followup_tags),
            'humanAppointmentFieldKeys', to_jsonb(auto_appointment_field_keys),
            'humanSaleFieldKeys', to_jsonb(auto_sale_field_keys),
            'scoreAttributeKey', to_jsonb(coalesce(suggested_score_key, ''))
        )
    );

    insert into cw.dashboard_tag_settings (account_id, settings, updated_at)
    values (target_account_id, refreshed_settings, now())
    on conflict (account_id)
    do update set settings = excluded.settings, updated_at = now();

    return jsonb_build_object(
        'account_id', target_account_id,
        'labels', label_list,
        'attribute_keys', all_attribute_keys,
        'autoTagGroups', refreshed_settings->'autoTagGroups',
        'updated_at', now()
    );
end;
$$;

revoke all on function cw.refresh_dashboard_discovery(bigint) from anon, authenticated;
grant execute on function cw.refresh_dashboard_discovery(bigint) to service_role;

select cw.refresh_dashboard_discovery(0);
select cw.prune_dashboard_settings_labels(0);
select cw.prune_deleted_label_references(0);

-- ========================================================================
-- FIN BLOQUE 12: addendum 2026-06-04 labels vivos, pruning y discovery
-- ========================================================================

-- ========================================================================
-- INICIO BLOQUE 13: limpieza opcional de conversaciones Chatwoot legacy
-- DONDE SE EJECUTA: Supabase > SQL Editor
-- INSTRUCCION: ejecutar solo si el proyecto fue clonado desde otra empresa
-- o si existen conversaciones con account_id de otro Chatwoot.
-- REEMPLAZAR antes de ejecutar:
--   <CHATWOOT_ACCOUNT_ID_ACTUAL> por el account_id real del negocio.
-- ========================================================================

begin;

drop table if exists _legacy_chatwoot_conversations;

create temporary table _legacy_chatwoot_conversations as
select c.chatwoot_conversation_id
from cw.conversations_current c
cross join lateral (
    select coalesce(
        c.chatwoot_account_id,
        case
            when c.raw_payload ->> 'account_id' ~ '^[0-9]+$'
            then (c.raw_payload ->> 'account_id')::bigint
        end
    ) as account_id
) account_scope
where account_scope.account_id is not null
  and account_scope.account_id <> <CHATWOOT_ACCOUNT_ID_ACTUAL>;

create table if not exists cw.archived_legacy_chatwoot_conversations as
select now()::timestamptz as archived_at,
       'legacy_chatwoot_account_mismatch'::text as archive_reason,
       c.*
from cw.conversations_current c
where false;

create table if not exists cw.archived_legacy_chatwoot_messages as
select now()::timestamptz as archived_at,
       'legacy_chatwoot_account_mismatch'::text as archive_reason,
       m.*
from cw.messages m
where false;

create table if not exists cw.archived_legacy_chatwoot_label_events as
select now()::timestamptz as archived_at,
       'legacy_chatwoot_account_mismatch'::text as archive_reason,
       e.*
from cw.conversation_label_events e
where false;

create table if not exists cw.archived_legacy_chatwoot_label_history as
select now()::timestamptz as archived_at,
       'legacy_chatwoot_account_mismatch'::text as archive_reason,
       h.*
from cw.conversation_label_history h
where false;

create table if not exists cw.archived_legacy_chatwoot_attribute_history as
select now()::timestamptz as archived_at,
       'legacy_chatwoot_account_mismatch'::text as archive_reason,
       h.*
from cw.conversation_attribute_history h
where false;

create table if not exists cw.archived_legacy_chatwoot_stage_history as
select now()::timestamptz as archived_at,
       'legacy_chatwoot_account_mismatch'::text as archive_reason,
       h.*
from cw.business_stage_history h
where false;

insert into cw.archived_legacy_chatwoot_conversations
select now(), 'legacy_chatwoot_account_mismatch', c.*
from cw.conversations_current c
join _legacy_chatwoot_conversations legacy using (chatwoot_conversation_id)
where not exists (
    select 1 from cw.archived_legacy_chatwoot_conversations archived
    where archived.chatwoot_conversation_id = c.chatwoot_conversation_id
);

insert into cw.archived_legacy_chatwoot_messages
select now(), 'legacy_chatwoot_account_mismatch', m.*
from cw.messages m
join _legacy_chatwoot_conversations legacy using (chatwoot_conversation_id)
where not exists (
    select 1 from cw.archived_legacy_chatwoot_messages archived
    where archived.chatwoot_conversation_id = m.chatwoot_conversation_id
      and archived.chatwoot_message_id is not distinct from m.chatwoot_message_id
);

insert into cw.archived_legacy_chatwoot_label_events
select now(), 'legacy_chatwoot_account_mismatch', e.*
from cw.conversation_label_events e
join _legacy_chatwoot_conversations legacy using (chatwoot_conversation_id)
where not exists (
    select 1 from cw.archived_legacy_chatwoot_label_events archived
    where archived.id = e.id
);

insert into cw.archived_legacy_chatwoot_label_history
select now(), 'legacy_chatwoot_account_mismatch', h.*
from cw.conversation_label_history h
join _legacy_chatwoot_conversations legacy using (chatwoot_conversation_id)
where not exists (
    select 1 from cw.archived_legacy_chatwoot_label_history archived
    where archived.id = h.id
);

insert into cw.archived_legacy_chatwoot_attribute_history
select now(), 'legacy_chatwoot_account_mismatch', h.*
from cw.conversation_attribute_history h
join _legacy_chatwoot_conversations legacy using (chatwoot_conversation_id)
where not exists (
    select 1 from cw.archived_legacy_chatwoot_attribute_history archived
    where archived.id = h.id
);

insert into cw.archived_legacy_chatwoot_stage_history
select now(), 'legacy_chatwoot_account_mismatch', h.*
from cw.business_stage_history h
join _legacy_chatwoot_conversations legacy using (chatwoot_conversation_id)
where not exists (
    select 1 from cw.archived_legacy_chatwoot_stage_history archived
    where archived.id = h.id
);

delete from cw.conversation_label_events e
using _legacy_chatwoot_conversations legacy
where e.chatwoot_conversation_id = legacy.chatwoot_conversation_id;

delete from cw.conversations_current c
using _legacy_chatwoot_conversations legacy
where c.chatwoot_conversation_id = legacy.chatwoot_conversation_id;

alter table cw.archived_legacy_chatwoot_conversations enable row level security;
alter table cw.archived_legacy_chatwoot_messages enable row level security;
alter table cw.archived_legacy_chatwoot_label_events enable row level security;
alter table cw.archived_legacy_chatwoot_label_history enable row level security;
alter table cw.archived_legacy_chatwoot_attribute_history enable row level security;
alter table cw.archived_legacy_chatwoot_stage_history enable row level security;

revoke all on table
    cw.archived_legacy_chatwoot_conversations,
    cw.archived_legacy_chatwoot_messages,
    cw.archived_legacy_chatwoot_label_events,
    cw.archived_legacy_chatwoot_label_history,
    cw.archived_legacy_chatwoot_attribute_history,
    cw.archived_legacy_chatwoot_stage_history
from anon, authenticated;

commit;

-- ========================================================================
-- FIN BLOQUE 13: limpieza opcional de conversaciones Chatwoot legacy
-- ========================================================================

-- ========================================================================
-- INICIO BLOQUE 14: validaciones finales addendum 2026-06-04
-- DONDE SE EJECUTA: Supabase > SQL Editor
-- INSTRUCCION: reemplazar <CHATWOOT_ACCOUNT_ID_ACTUAL> antes de ejecutar.
-- ========================================================================

select jobid, jobname, schedule, active,
       position('/functions/v1/chatwoot-sync' in command) > 0 as invokes_chatwoot_sync,
       position('mode' in command) > 0 as sends_body,
       position('window_hours' in command) > 0 as sends_window
from cron.job
where jobname = 'sync-chatwoot-diario';

select id, sync_type, status, started_at, finished_at,
       stats ->> 'mode' as mode,
       stats ->> 'conversations' as conversations,
       stats ->> 'messages' as messages,
       stats ->> 'labels' as labels,
       left(coalesce(error_message, ''), 240) as error_message
from cw.sync_runs
order by started_at desc nulls last, id desc
limit 8;

select cursor_name, last_since_ts, last_until_ts, updated_at, cursor_payload
from cw.sync_cursor
order by updated_at desc;

select
    (select count(*) from cw.inboxes) as inboxes,
    (select count(*) from cw.label_catalog) as labels,
    (select count(*) from cw.attribute_definitions) as attribute_definitions,
    (select count(*) from cw.contacts_current) as contacts,
    (select count(*) from cw.conversations_current where chatwoot_account_id = <CHATWOOT_ACCOUNT_ID_ACTUAL>) as current_chatwoot_conversations,
    (select count(*) from cw.messages where chatwoot_account_id = <CHATWOOT_ACCOUNT_ID_ACTUAL>) as current_chatwoot_messages,
    (select count(*) from cw.conversations_current where source_system = 'n8n_chat_histories') as historical_n8n_conversations,
    (select count(*) from cw.messages m join cw.conversations_current c using (chatwoot_conversation_id) where c.source_system = 'n8n_chat_histories') as historical_n8n_messages,
    (select count(*)
     from cw.conversations_current c
     where coalesce(
        c.chatwoot_account_id,
        case
            when c.raw_payload ->> 'account_id' ~ '^[0-9]+$'
            then (c.raw_payload ->> 'account_id')::bigint
        end
     ) is not null
       and coalesce(
        c.chatwoot_account_id,
        case
            when c.raw_payload ->> 'account_id' ~ '^[0-9]+$'
            then (c.raw_payload ->> 'account_id')::bigint
        end
       ) <> <CHATWOOT_ACCOUNT_ID_ACTUAL>) as active_legacy_conversations;

with active_labels as (
    select title
    from cw.label_catalog
    where account_id = 0
),
conversation_stale as (
    select distinct label
    from cw.conversations_current c,
         unnest(coalesce(c.labels, array[]::text[])) as labels(label)
    where not exists (select 1 from active_labels a where a.title = label)
),
event_stale as (
    select distinct label
    from cw.conversation_label_events e,
         unnest(
            coalesce(e.previous_labels, array[]::text[]) ||
            coalesce(e.next_labels, array[]::text[]) ||
            coalesce(e.added_labels, array[]::text[]) ||
            coalesce(e.removed_labels, array[]::text[])
         ) as labels(label)
    where not exists (select 1 from active_labels a where a.title = label)
)
select
    (select count(*) from conversation_stale) as stale_conversation_labels,
    (select count(*) from event_stale) as stale_event_labels;

-- ========================================================================
-- FIN BLOQUE 14: validaciones finales addendum 2026-06-04
-- ========================================================================

-- ========================================================================
-- INICIO BLOQUE 15: runbook de Edge Functions y secrets por negocio
-- ESTE BLOQUE ES DOCUMENTAL; NO EJECUTAR COMO SQL.
-- ========================================================================
--
-- Secrets minimos por negocio:
--   SUPABASE_URL
--   SUPABASE_SERVICE_ROLE_KEY
--   VITE_CHATWOOT_BASE_URL / CHATWOOT_BASE_URL
--   VITE_CHATWOOT_ACCOUNT_ID / CHATWOOT_ACCOUNT_ID
--   VITE_CHATWOOT_API_TOKEN / CHATWOOT_API_TOKEN
--   CHATWOOT_WEBHOOK_SECRET
--   OPENAI_API_KEY
--   OPENAI_REPORT_MODEL
--   OPENAI_REPORT_REASONING_EFFORT
--   RESEND_API_KEY
--   RESEND_FROM_EMAIL
--   ZAPAWAY_API_KEY o proveedor WhatsApp equivalente, si aplica.
--
-- Deploy:
--   npx supabase functions deploy chatwoot-sync chatwoot-repair-conversations generate-ai-report meta-campaign-insights send-scheduled-reports --project-ref <PROJECT_REF> --use-api --jobs 2
--   npx supabase functions deploy chatwoot-label-webhook --project-ref <PROJECT_REF> --use-api --no-verify-jwt
--
-- Webhook Chatwoot:
--   https://<PROJECT_REF>.supabase.co/functions/v1/chatwoot-label-webhook?secret=<CHATWOOT_WEBHOOK_SECRET>
--
-- Sync manual:
--   POST https://<PROJECT_REF>.supabase.co/functions/v1/chatwoot-sync
--   Body: {"mode":"full","window_hours":72,"sync_messages":"recent","dashboard_account_id":0}
--
-- Criterio de cierre:
--   chatwoot-sync ACTIVE, cron active, sync_runs success, labels sin stale,
--   active_legacy_conversations = 0 y npm run check aprobado.
--
-- ========================================================================
-- FIN BLOQUE 15: runbook de Edge Functions y secrets por negocio
-- ========================================================================
"""


SQL_META_ADS_ADDENDUM = r"""

-- ========================================================================
-- INICIO BLOQUE 16: Meta Ads configurable por negocio
-- DONDE SE EJECUTA: Supabase > SQL Editor
-- INSTRUCCION: ejecutar despues del esquema base de cw y antes de usar
-- Tendencias > Meta Ads > Configurar campanas.
-- OBJETIVO: guardar configuracion Meta Ads server-side sin exponer el
-- Bearer Token al frontend. No insertar tokens reales en este SQL.
-- ========================================================================

create schema if not exists cw;
create extension if not exists pgcrypto with schema extensions;

create table if not exists cw.meta_ads_configs (
    id uuid primary key default gen_random_uuid(),
    account_id bigint not null default 0,
    ad_account_id text not null,
    access_token text not null,
    token_last_four text not null default '',
    graph_api_version text not null default 'v20.0',
    enabled boolean not null default true,
    configured_by uuid,
    configured_at timestamptz not null default now(),
    created_at timestamptz not null default now(),
    updated_at timestamptz not null default now(),
    unique (account_id)
);

create index if not exists meta_ads_configs_account_enabled_idx
    on cw.meta_ads_configs (account_id, enabled);

alter table cw.meta_ads_configs enable row level security;

revoke all on cw.meta_ads_configs from anon, authenticated;

grant usage on schema cw to service_role;
grant select, insert, update, delete on cw.meta_ads_configs to service_role;

-- Validacion de seguridad esperada:
--   relrowsecurity = true
--   sin grants para anon/authenticated
--   grants de lectura/escritura solo para service_role/postgres.
select c.relname,
       c.relrowsecurity,
       coalesce(
           jsonb_agg(
               jsonb_build_object('grantee', g.grantee, 'privilege', g.privilege_type)
               order by g.grantee, g.privilege_type
           ) filter (where g.grantee is not null),
           '[]'::jsonb
       ) as grants
from pg_class c
join pg_namespace n on n.oid = c.relnamespace
left join information_schema.role_table_grants g
       on g.table_schema = n.nspname
      and g.table_name = c.relname
where n.nspname = 'cw'
  and c.relname = 'meta_ads_configs'
group by c.relname, c.relrowsecurity;

-- Runbook de configuracion:
-- 1. Desplegar Edge Function:
--    npx supabase functions deploy meta-campaign-insights --project-ref <PROJECT_REF> --use-api
-- 2. Ingresar al dashboard con rol platform_admin o company_admin.
-- 3. Abrir Tendencias > Meta Ads > Configurar campanas.
-- 4. Guardar Meta Ad Account ID, Bearer Token y Graph API Version opcional.
-- 5. Pulsar Actualizar para cargar campaigns e insights.
--
-- La funcion meta-campaign-insights soporta:
--   action = get_config     -> devuelve configured, adAccountId, tokenLast4 y graphApiVersion.
--   action = save_config    -> guarda/rota token, solo para platform_admin/company_admin.
--   action omitida/fetch    -> trae campaigns e insights usando cw.meta_ads_configs.
--
-- Fallback legacy opcional por Supabase Secrets:
--   META_AD_ACCOUNT_ID
--   META_SYSTEM_USER_TOKEN
--   META_GRAPH_API_VERSION
--   META_CACHE_TTL_SECONDS
-- Estos secretos solo se usan si cw.meta_ads_configs aun no tiene fila activa.

-- Validacion posterior a configurar desde el dashboard:
select account_id,
       ad_account_id,
       token_last_four,
       graph_api_version,
       enabled,
       configured_at,
       updated_at
from cw.meta_ads_configs
order by updated_at desc;

select ad_account_id,
       date_start,
       date_stop,
       status,
       started_at,
       finished_at,
       campaign_rows,
       insight_rows,
       returned_rows,
       left(coalesce(error_message, ''), 240) as error_message
from cw.meta_ads_sync_runs
order by started_at desc
limit 10;

-- ========================================================================
-- FIN BLOQUE 16: Meta Ads configurable por negocio
-- ========================================================================
"""


def update_sql() -> bool:
    text = SQL_PATH.read_text(encoding="utf-8")
    changed = False

    updated_header = text.replace(
        "-- Fecha base: 2026-05-18. Actualización documental: 2026-05-29.",
        "-- Fecha base: 2026-05-18. Actualización documental: 2026-06-04.",
    )
    if updated_header != text:
        text = updated_header
        changed = True

    if SQL_MARKER not in text:
        text = text.rstrip() + SQL_ADDENDUM + "\n"
        changed = True

    if SQL_META_ADS_MARKER not in text:
        text = text.rstrip() + SQL_META_ADS_ADDENDUM + "\n"
        changed = True

    if changed:
        with SQL_PATH.open("w", encoding="utf-8", newline="\n") as handle:
            handle.write(text)

    return changed


def mirror_requested_outputs() -> list[Path]:
    mirrored: list[Path] = []

    REQUESTED_DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    REQUESTED_SQL_PATH.parent.mkdir(parents=True, exist_ok=True)

    if DOCX_PATH.resolve() != REQUESTED_DOCX_PATH.resolve():
        if REQUESTED_DOCX_PATH.exists():
            backup(REQUESTED_DOCX_PATH)
        REQUESTED_DOCX_PATH.write_bytes(DOCX_PATH.read_bytes())
        mirrored.append(REQUESTED_DOCX_PATH)

    if SQL_PATH.resolve() != REQUESTED_SQL_PATH.resolve():
        if REQUESTED_SQL_PATH.exists():
            backup(REQUESTED_SQL_PATH)
        REQUESTED_SQL_PATH.write_bytes(SQL_PATH.read_bytes())
        mirrored.append(REQUESTED_SQL_PATH)

    return mirrored


def main() -> None:
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)
    if not SQL_PATH.exists():
        raise FileNotFoundError(SQL_PATH)

    docx_backup = backup(DOCX_PATH)
    sql_backup = backup(SQL_PATH)

    docx_changed = update_docx()
    docx_meta_ads_changed = update_docx_meta_ads()
    sql_changed = update_sql()
    mirrored_paths = mirror_requested_outputs()

    print(f"DOCX source: {DOCX_PATH}")
    print(f"SQL source: {SQL_PATH}")
    print(f"DOCX backup: {docx_backup}")
    print(f"SQL backup: {sql_backup}")
    print(f"DOCX changed: {docx_changed or docx_meta_ads_changed}")
    print(f"DOCX v2.2 changed: {docx_changed}")
    print(f"DOCX Meta Ads v2.3 changed: {docx_meta_ads_changed}")
    print(f"SQL changed: {sql_changed}")
    print("Mirrored outputs:")
    for mirrored_path in mirrored_paths:
        print(f"- {mirrored_path}")


if __name__ == "__main__":
    main()
