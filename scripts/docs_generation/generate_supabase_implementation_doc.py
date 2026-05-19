from __future__ import annotations

from datetime import date
from pathlib import Path
from textwrap import dedent

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[2]
DOCS_DIR = REPO_ROOT / "docs"
OUTPUT_DOCX = DOCS_DIR / "Implementacion_Supabase_SimpliaLeads.docx"
HISTORICAL_DOCX = DOCS_DIR / "Implementacion_Supabase_SimpliaLeads_ISO10013_v1_2.docx"
LEGACY_SQL = DOCS_DIR / "supabase_replicacion_simpliale_solo_public_cw.sql"
MIGRATIONS_DIR = REPO_ROOT / "supabase" / "migrations"

BASE_SQL_BLOCKS = [
    ["01", "Public mínimo + seguridad", "Crea public.user_profiles, tipo public.app_role, trigger de auth.users y RLS básico. La tabla public.dashboard_tag_settings aparece en el SQL legado, pero la migración final la elimina."],
    ["02", "Schema cw base", "Crea el schema cw y sus tablas principales para Chatwoot: configuración, sync, contactos, conversaciones, mensajes, métricas, reportes y auditoría."],
    ["03", "Funciones, triggers e índices cw", "Agrega updated_at automático, captura de historial de atributos/etiquetas, índices de rendimiento y políticas RLS/grants."],
    ["04", "Vault para cron", "Crea secrets internos de Vault para que pg_cron invoque Edge Functions usando pg_net."],
    ["05", "Migraciones locales acumuladas", "Agrega vistas de salud, cron diario, reportes programados, canales, auditoría comercial e importaciones."],
    ["06", "Seeds/configuración inicial", "Inserta configuración inicial de cw.dashboard_tag_settings. En la versión final esta configuración es editable y no debe depender de public."],
    ["07", "Roles de usuarios", "Actualiza public.user_profiles después de crear usuarios manualmente en Authentication."],
    ["08", "Validación", "Consultas para comprobar extensiones, tablas, roles, cron, Vault y estado de sync."],
]

CW_TABLE_DICTIONARY = [
    ["cw.account_config", "Configuración de cuenta Chatwoot/Simplia para el sync.", "SQL base", "Manual/seed; puede ser usada por procesos server-side."],
    ["cw.sync_cursor", "Cursor de sincronización incremental.", "chatwoot-sync", "Permite recordar hasta dónde se sincronizó."],
    ["cw.sync_runs", "Bitácora de cada ejecución de sync.", "chatwoot-sync / cron / manual", "Clave para diagnosticar si el cron corrió y si terminó success/error."],
    ["cw.raw_ingest", "Payloads crudos recibidos o importados.", "sync/importaciones", "Soporte y auditoría cuando un dato normalizado no cuadra."],
    ["cw.inboxes", "Inboxes/canales de Chatwoot.", "chatwoot-sync", "Permite filtrar por canal e identificar WhatsApp, Instagram, web, etc."],
    ["cw.teams", "Equipos de Chatwoot si la cuenta los expone.", "chatwoot-sync", "Base para reporting operativo por equipo."],
    ["cw.attribute_definitions", "Definiciones oficiales de custom attributes de Chatwoot.", "chatwoot-sync", "Automático; no se escribe manualmente. Alimenta opciones del dashboard."],
    ["cw.contacts_current", "Snapshot actual de contactos.", "chatwoot-sync / repair", "Guarda custom_attributes de contacto y datos identificadores."],
    ["cw.contact_inboxes", "Relación contacto-inbox.", "chatwoot-sync", "Útil para saber por qué canal llegó un contacto."],
    ["cw.conversations_current", "Snapshot actual de conversaciones.", "chatwoot-sync / webhook / repair / import", "Tabla principal del dashboard histórico y KPIs."],
    ["cw.contact_attribute_history", "Historial de cambios en atributos de contacto.", "triggers de sync/repair", "Auditoría de cambios de campos personalizados."],
    ["cw.conversation_attribute_history", "Historial de cambios en atributos de conversación.", "triggers de sync/repair", "Permite ver evolución de score, monto, fecha, responsable, etc."],
    ["cw.conversation_label_history", "Historial simple de etiquetas agregadas/removidas.", "sync/webhook", "Auditoría de cambios comerciales por etiqueta."],
    ["cw.business_stage_history", "Historial de etapa comercial resuelta.", "triggers/configuración", "Ayuda a entender cambios de embudo."],
    ["cw.messages", "Mensajes normalizados por conversación.", "chatwoot-sync", "Base para reportes de conversaciones, tráfico entrante y tiempos."],
    ["cw.reporting_events", "Eventos de reporting de Chatwoot.", "chatwoot-sync", "Métricas operativas cuando Chatwoot expone reporting_events."],
    ["cw.daily_metrics", "Métricas agregadas por día.", "sync/procesos de reporting", "Base para tendencias y resúmenes."],
    ["cw.automated_reports", "Configuración de reportes programados.", "dashboard", "Define destinatarios, frecuencia, formato, filtros y estado activo."],
    ["cw.dashboard_tag_settings", "Configuración humana del dashboard.", "dashboard/admin", "Única fuente de verdad para SQL/cita/venta/no calificado/campos IA/contexto empresarial."],
    ["cw.conversation_label_events", "Eventos deduplicados de cambios de etiquetas.", "webhook/sync diff", "Permite detectar cambios relevantes sin duplicarlos."],
    ["cw.automated_report_runs", "Resultado de cada ejecución de reporte programado.", "send-scheduled-reports", "Evidencia de envíos, errores y auditoría."],
    ["cw.import_batches", "Lotes de importación manual/flexible.", "importador", "Permite rastrear cargas masivas externas."],
    ["cw.import_batch_errors", "Errores por lote de importación.", "importador", "Diagnóstico de filas rechazadas o incompletas."],
    ["cw.label_catalog", "Catálogo automático de etiquetas reales de Chatwoot.", "chatwoot-sync / webhook", "Se descubre solo; no decide significado de negocio."],
    ["cw.attribute_key_catalog", "Catálogo automático de llaves de atributos vistas en definiciones y payloads.", "chatwoot-sync / refresh_dashboard_discovery", "Ayuda al dashboard a mostrar campos disponibles aunque cambien por negocio."],
]

PUBLIC_TABLE_DICTIONARY = [
    ["public.user_profiles", "Perfil de usuario vinculado a auth.users.", "Trigger on_auth_user_created", "Se crea después de crear manualmente usuarios en Authentication."],
]

MIGRATION_DICTIONARY = [
    ["20260421134500_create_conversation_label_events.sql", "Crea cw.conversation_label_events e índices GIN para added/removed labels.", "Ejecutar después del schema cw base si no está incluido en el SQL consolidado."],
    ["20260426190000_harden_chatwoot_snapshots_and_sync.sql", "Endurece snapshots, índices, RLS, vista cw.sync_health, Vault y cron diario.", "Crítico para operación y diagnóstico."],
    ["20260426202500_grant_cw_to_service_role.sql", "Asegura permisos de service_role sobre cw.", "Necesario para Edge Functions con service role."],
    ["20260427165000_reporting_exports_and_scheduled_runs.sql", "Agrega reportes programados, runs y cron send-scheduled-reports.", "Necesario para correos programados."],
    ["20260427220500_backfill_chatwoot_inbox_channel.sql", "Rellena canales/inbox históricos.", "Normaliza canal en datos existentes."],
    ["20260427225200_enforce_chatwoot_inbox_channel.sql", "Función/trigger para resolver canal Chatwoot de forma consistente.", "Mantiene canal estable en conversaciones."],
    ["20260429123000_commercial_current_state_audit.sql", "Auditoría comercial de estado actual.", "Base para cambios relevantes y reportes."],
    ["20260504190000_flexible_lead_import.sql", "Lotes y errores para importación flexible de leads.", "Permite cargar datos externos con trazabilidad."],
    ["20260519130831_auto_discover_chatwoot_catalogs.sql", "Crea label_catalog, attribute_key_catalog y helpers de discovery.", "Hace automático el descubrimiento de etiquetas/atributos."],
    ["20260519133145_centralize_dashboard_settings_in_cw.sql", "Centraliza dashboard settings en cw y elimina public.dashboard_tag_settings.", "Debe ejecutarse al final para dejar el modelo vigente."],
]

EDGE_FUNCTION_DICTIONARY = [
    ["chatwoot-sync", "Sincroniza Chatwoot hacia cw: inboxes, labels, attribute_definitions, contacts, conversations, messages y catálogos.", "JWT activo", "Manual, cron diario 00:01 Guayaquil, reparación inicial."],
    ["chatwoot-repair-conversations", "Repara o refresca conversaciones específicas desde Chatwoot.", "JWT activo", "Uso admin/soporte cuando un dato puntual no cuadra."],
    ["chatwoot-label-webhook", "Recibe eventos de Chatwoot y actualiza snapshots/eventos/catálogos.", "JWT desactivado + CHATWOOT_WEBHOOK_SECRET", "Live operacional por webhook; Chatwoot no envía JWT Supabase."],
    ["generate-ai-report", "Genera reportes IA bajo demanda usando OpenAI.", "JWT activo", "Dashboard > Reportes."],
    ["send-scheduled-reports", "Busca reportes activos, genera contenido y envía correos con Resend.", "JWT activo", "Cron cada 5 minutos."],
]

COLORS = {
    "navy": "1f3f91",
    "blue": "274690",
    "ink": "0f2344",
    "slate": "64748b",
    "line": "d9e2ef",
    "bg": "f8fafc",
    "white": "ffffff",
    "green": "0a9b6f",
    "mint": "dff7ec",
    "orange": "f59e0b",
    "yellow": "fff7ed",
}


def rgb(hex_color: str) -> RGBColor:
    value = hex_color.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text: str, *, bold: bool = False, fill: str | None = None, color: str | None = None, size: float = 8.0) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if fill:
        shade(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def style_table(table, header_fill: str = COLORS["blue"]) -> None:
    table.style = "Table Grid"
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(7.6 if row_index else 8.0)
            if row_index == 0:
                shade(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = rgb(COLORS["white"])
            else:
                shade(cell, COLORS["white"] if row_index % 2 else COLORS["bg"])


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)

    for name, size, color in [
        ("Title", 20, COLORS["blue"]),
        ("Heading 1", 15, COLORS["blue"]),
        ("Heading 2", 12, COLORS["ink"]),
        ("Heading 3", 10.5, COLORS["ink"]),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("IMP-SUPA-SIMPLIA-002 | Version 2.0 | Documento controlado | Pagina ")
    run.font.name = "Arial"
    run.font.size = Pt(8)


def add_title(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = paragraph.add_run("Implementación Supabase SimpliaLeads")
    title.bold = True
    title.font.name = "Arial"
    title.font.size = Pt(22)
    title.font.color.rgb = rgb(COLORS["blue"])

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Runbook operativo para desarrolladores e IA: base de datos, Edge Functions, secrets, Chatwoot, cron y pruebas")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = rgb(COLORS["slate"])


def h(document: Document, level: int, text: str) -> None:
    document.add_heading(text, level=level)


def p(document: Document, text: str) -> None:
    document.add_paragraph(text)


def bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def numbered(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


def table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    tbl = document.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        set_cell(tbl.rows[0].cells[index], header, bold=True, fill=COLORS["blue"], color=COLORS["white"], size=8.0)
    for row in rows:
        cells = tbl.add_row().cells
        for index, value in enumerate(row):
            set_cell(cells[index], value, size=7.6)
    if widths:
        for row in tbl.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    style_table(tbl)
    document.add_paragraph()


def code(document: Document, text: str, title: str | None = None) -> None:
    if title:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(title)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(8.5)
        run.font.color.rgb = rgb(COLORS["ink"])

    clean = dedent(text).strip("\n")
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    for line in clean.splitlines():
        run = paragraph.add_run(line.rstrip() + "\n")
        run.font.name = "Consolas"
        run.font.size = Pt(7.0)
        run.font.color.rgb = rgb(COLORS["ink"])


def callout(document: Document, title: str, body: str, tone: str = "blue") -> None:
    fill = COLORS["mint"] if tone == "green" else COLORS["yellow"] if tone == "yellow" else "eef4ff"
    tbl = document.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    shade(cell, fill)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = rgb(COLORS["ink"])
    paragraph.add_run("\n")
    body_run = paragraph.add_run(body)
    body_run.font.name = "Arial"
    body_run.font.size = Pt(8)
    body_run.font.color.rgb = rgb(COLORS["ink"])
    document.add_paragraph()


def read_file(path: Path, limit: int | None = None) -> str:
    if not path.exists():
        return f"-- Archivo no encontrado: {path.relative_to(REPO_ROOT)}"
    value = path.read_text(encoding="utf-8")
    if limit and len(value) > limit:
        return value[:limit] + "\n\n-- [TRUNCADO EN DOCUMENTO: revisar archivo fuente completo en el repo]\n"
    return value


def build_document() -> None:
    document = Document()
    configure_document(document)
    add_title(document)

    h(document, 1, "1. Control Documental")
    table(
        document,
        ["Campo", "Valor"],
        [
            ["Código", "IMP-SUPA-SIMPLIA-002"],
            ["Versión", "2.0"],
            ["Fecha", date.today().strftime("%d/%m/%Y")],
            ["Documento base histórico", str(HISTORICAL_DOCX.relative_to(REPO_ROOT))],
            ["Documento generado", str(OUTPUT_DOCX.relative_to(REPO_ROOT))],
            ["Modelo documental", "ISO 10013 / información documentada para implementación y control operacional."],
            ["Alcance", "Proyecto Supabase por negocio para SimpliaLeads: schema cw, roles, Edge Functions, secrets, Vault, cron, Chatwoot y pruebas."],
            ["Exclusiones", "No cubre el rediseño SaaS multiempresa pooled con company_id. No almacena secretos reales."],
        ],
        [1.9, 5.3],
    )

    h(document, 1, "2. Resumen Ejecutivo")
    p(
        document,
        "Este documento es el runbook oficial para replicar o reconstruir la implementación Supabase de SimpliaLeads. "
        "Está escrito para que lo pueda seguir un desarrollador humano o una IA con acceso al repositorio, Supabase CLI y permisos Owner/Admin."
    )
    callout(
        document,
        "Regla central de la versión 2.0",
        "La configuración vigente del dashboard vive solo en cw.dashboard_tag_settings. "
        "La tabla public.dashboard_tag_settings no debe quedar como fuente de verdad; si existe por un SQL legado, debe eliminarse con la migración de centralización.",
        "green",
    )
    bullets(
        document,
        [
            "Chatwoot es la fuente de verdad de conversaciones, etiquetas, contactos, inboxes y atributos personalizados.",
            "Supabase guarda histórico, snapshots actuales, catálogos detectados, auditoría, reportes y estado de sincronización.",
            "El dashboard combina datos live desde Chatwoot con histórico desde Supabase.",
            "El sistema detecta etiquetas y atributos automáticamente; la interpretación de negocio se configura manualmente en el dashboard.",
        ],
    )

    h(document, 1, "3. Modelo Final Vigente")
    table(
        document,
        ["Objeto", "Estado final", "Uso"],
        [
            ["cw.dashboard_tag_settings", "Vigente y único", "Configuración humana del dashboard: SQL, citas, ventas, no calificados, campos de score, contexto empresarial."],
            ["cw.label_catalog", "Automático", "Catálogo de etiquetas reales detectadas desde Chatwoot por sync/webhook."],
            ["cw.attribute_definitions", "Automático", "Definiciones oficiales de custom attributes desde Chatwoot."],
            ["cw.attribute_key_catalog", "Automático", "Llaves detectadas en definiciones y payloads reales de contactos/conversaciones."],
            ["cw.contacts_current", "Automático", "Snapshot actual de contactos con custom_attributes."],
            ["cw.conversations_current", "Automático", "Snapshot actual de conversaciones, etiquetas, atributos y datos comerciales."],
            ["cw.messages", "Automático", "Mensajes normalizados para histórico y reportes."],
            ["cw.sync_runs", "Automático", "Bitácora de cada corrida de sync."],
            ["public.user_profiles", "Vigente", "Perfiles de usuarios Auth y roles de app."],
            ["public.dashboard_tag_settings", "No vigente", "No debe usarse como fuente de verdad. La migración final lo elimina."],
        ],
        [1.9, 1.4, 4.0],
    )
    p(
        document,
        "La diferencia importante es intención: descubrir no significa decidir. "
        "Supabase puede saber que existe una etiqueta llamada venta_exitosa, pero no debe decidir por sí solo que esa etiqueta cuenta como venta; esa decisión queda en cw.dashboard_tag_settings."
    )

    h(document, 2, "3.1 Decisiones De Arquitectura")
    table(
        document,
        ["Decisión", "Por qué se hace así", "Qué problema evita"],
        [
            [
                "Usar schema cw y no public para datos de Chatwoot",
                "cw separa el dominio Chatwoot/dashboard del schema public, que en Supabase suele estar expuesto por la Data API. "
                "Esto deja public mínimo y reduce el riesgo de exponer tablas operativas por accidente.",
                "Mezclar datos sensibles, histórico, mensajes, contactos y reportes en public; fuga accidental por grants/RLS mal configurados.",
            ],
            [
                "Mantener public solo para user_profiles",
                "user_profiles está ligado a auth.users y al rol de aplicación. Es pequeño, transversal y necesita integrarse con Auth/RLS.",
                "Duplicar usuarios en cw o guardar roles dentro de payloads de Chatwoot.",
            ],
            [
                "Guardar configuración del dashboard en cw.dashboard_tag_settings",
                "Las reglas de negocio del dashboard dependen de etiquetas/atributos Chatwoot. Por eso pertenecen al dominio cw y no a public.",
                "Tener dos fuentes de verdad como public.dashboard_tag_settings y cw.dashboard_tag_settings.",
            ],
            [
                "Usar Edge Function Secrets y no .env para claves privadas",
                ".env del frontend puede terminar en builds, previews, logs o máquinas locales. Las Edge Functions corren server-side y leen secrets sin exponerlos al navegador.",
                "Exponer OPENAI_API_KEY, RESEND_API_KEY, service role key o Chatwoot API token en cliente o repositorio.",
            ],
            [
                "Usar Vault para pg_cron",
                "pg_cron corre dentro de Postgres y necesita llamar Edge Functions con pg_net. Vault guarda URL/JWT dentro de la base sin hardcodearlos en jobs.",
                "Dejar JWTs escritos directamente dentro del SQL de cron o scripts compartidos.",
            ],
            [
                "Desactivar JWT solo en chatwoot-label-webhook",
                "Chatwoot no envía JWT Supabase. Ese endpoint se protege con CHATWOOT_WEBHOOK_SECRET compartido.",
                "Webhook roto por 401 o endpoint público sin validación de secreto.",
            ],
        ],
        [1.9, 3.0, 2.5],
    )
    callout(
        document,
        "Regla de seguridad para desarrolladores",
        "Todo secreto que permita escribir, sincronizar, enviar correos o consumir OpenAI debe vivir server-side: Supabase Edge Function Secrets o Vault. "
        "El .env del frontend solo debe contener valores públicos o temporales de desarrollo, nunca service role ni API keys privadas.",
        "yellow",
    )

    h(document, 2, "3.2 Diccionario Funcional De Tablas")
    p(
        document,
        "Este diccionario permite que un desarrollador nuevo entienda para qué existe cada tabla antes de ejecutar o modificar SQL. "
        "Las tablas de cw pertenecen al dominio Chatwoot/dashboard; public se mantiene mínimo por seguridad."
    )
    table(
        document,
        ["Tabla", "Qué guarda", "Quién la llena", "Por qué existe"],
        PUBLIC_TABLE_DICTIONARY + CW_TABLE_DICTIONARY,
        [1.8, 2.1, 1.7, 2.4],
    )

    h(document, 1, "4. Requisitos Previos")
    table(
        document,
        ["Requisito", "Cómo se obtiene", "Validación"],
        [
            ["Repositorio", "Clonar o abrir este repo en la máquina de trabajo.", "Existe carpeta supabase/functions y docs."],
            ["Node.js/npm", "Instalación local para ejecutar npx y build.", "npm --version"],
            ["Supabase CLI", "Usar npx supabase o CLI instalada.", "npx supabase --version"],
            ["Token Supabase", "Supabase Dashboard > Account > Access Tokens.", "npx supabase login --token \"<SUPABASE_ACCESS_TOKEN>\""],
            ["Permisos Supabase", "Owner/Admin del proyecto destino.", "Puede ver API keys, SQL Editor, Edge Functions, Secrets, Database Extensions."],
            ["Chatwoot", "Base URL, account ID y API token de administrador.", "API /inboxes, /labels y /custom_attribute_definitions responden."],
            ["OpenAI", "API key de proyecto.", "Se guarda como secret; no se imprime."],
            ["Resend", "API key y dominio verificado.", "RESEND_FROM_EMAIL usa dominio verificado."],
        ],
        [1.6, 3.0, 2.5],
    )

    h(document, 1, "5. Datos Que Debe Entregar El Usuario")
    table(
        document,
        ["Dato", "Placeholder", "Notas"],
        [
            ["Supabase project ref", "<PROJECT_REF>", "Ejemplo visual: knsmqbkdsfhttizaepzv. No inventar."],
            ["Supabase URL", "https://<PROJECT_REF>.supabase.co", "Se usa en frontend, Edge Functions y Vault."],
            ["Supabase anon key", "<SUPABASE_ANON_KEY>", "Para frontend e invocaciones autenticadas con JWT activo."],
            ["Supabase service role key", "<SUPABASE_SERVICE_ROLE_KEY>", "Solo backend/Vault/Edge Functions. Nunca en frontend."],
            ["Supabase access token", "<SUPABASE_ACCESS_TOKEN>", "Token personal para CLI; no se guarda en docs."],
            ["Chatwoot base URL", "<CHATWOOT_BASE_URL>", "Ejemplo: https://app.chatwoot.com. Sin /api/v1/accounts."],
            ["Chatwoot account ID", "<CHATWOOT_ACCOUNT_ID>", "ID numérico de la cuenta."],
            ["Chatwoot API token", "<CHATWOOT_API_TOKEN>", "Token de usuario admin/owner con permisos de lectura."],
            ["OpenAI API key", "<OPENAI_API_KEY>", "Entregado por el usuario; se guarda como secret."],
            ["Resend API key", "<RESEND_API_KEY>", "Entregado por el usuario; se guarda como secret."],
            ["Email remitente Resend", "Simplia Leads <reportes@simpliaconsulting.com>", "Debe pertenecer a un dominio verificado."],
            ["Webhook secret", "<CHATWOOT_WEBHOOK_SECRET>", "Generado aleatoriamente; mismo valor en Supabase y Chatwoot."],
        ],
        [1.7, 2.4, 3.1],
    )

    h(document, 1, "6. Validar .env Antes De Empezar")
    p(document, "Antes de tocar Supabase, confirmar que el archivo .env del repo apunta al Chatwoot y Supabase correctos.")
    table(
        document,
        ["Variable .env", "Origen", "Regla"],
        [
            ["VITE_SUPABASE_URL", "Supabase API URL", "Debe ser https://<PROJECT_REF>.supabase.co."],
            ["VITE_SUPABASE_ANON_KEY", "Supabase anon key", "Key pública para frontend."],
            ["VITE_CHATWOOT_BASE_URL", "Chatwoot", "Debe ser solo la URL base, por ejemplo https://app.chatwoot.com."],
            ["VITE_CHATWOOT_ACCOUNT_ID", "Chatwoot", "Debe coincidir con la cuenta donde viven conversaciones, etiquetas e inboxes."],
            ["VITE_CHATWOOT_API_TOKEN", "Chatwoot", "Token usado por el frontend/dev y como base para secrets de Edge Functions."],
        ],
        [2.2, 1.8, 3.2],
    )
    code(
        document,
        r"""
        # PowerShell: listar nombres sin exponer valores
        Get-Content .env |
          Where-Object { $_ -match '^\s*[^#][^=]+=' } |
          ForEach-Object { ($_ -split '=', 2)[0].Trim() } |
          Sort-Object
        """,
        "Validación local de .env",
    )

    h(document, 1, "7. Orden Exacto De Implementación")
    numbered(
        document,
        [
            "Crear o seleccionar el proyecto Supabase destino y obtener PROJECT_REF, URL, anon key y service role key.",
            "Validar .env con los valores correctos de Supabase y Chatwoot.",
            "Iniciar sesión en Supabase CLI con token personal.",
            "Vincular el repo al proyecto o usar siempre --project-ref.",
            "Habilitar/verificar extensiones pgcrypto, pg_net, supabase_vault y pg_cron.",
            "Configurar timezone de base con America/Guayaquil.",
            "Aplicar SQL base/migraciones y luego las correcciones finales de discovery y centralización en cw.",
            "Crear usuarios en Authentication y asignar roles en public.user_profiles.",
            "Configurar Edge Function Secrets con placeholders reales entregados por el usuario.",
            "Configurar Vault secrets para que pg_cron invoque Edge Functions.",
            "Desplegar Edge Functions con JWT correcto según cada caso.",
            "Exponer schema cw en Data API si el dashboard lo requiere.",
            "Crear webhook en Chatwoot con CHATWOOT_WEBHOOK_SECRET.",
            "Ejecutar sync inicial de 72h y pruebas de reportes.",
            "Validar cron, logs, tablas, catálogos y limpiar datos temporales de prueba.",
        ],
    )

    h(document, 1, "8. Supabase CLI")
    code(
        document,
        r"""
        npx supabase --version
        npx supabase login --token "<SUPABASE_ACCESS_TOKEN>"
        npx supabase link --project-ref "<PROJECT_REF>"
        npx supabase secrets list --project-ref "<PROJECT_REF>"
        """,
        "Comandos iniciales",
    )
    callout(
        document,
        "Para IA/agentes",
        "Si falta un valor sensible, la IA debe pedirlo y no inventarlo. Nunca debe imprimir OPENAI_API_KEY, RESEND_API_KEY, CHATWOOT_API_TOKEN ni service role key en la respuesta final.",
        "yellow",
    )

    h(document, 1, "9. SQL, Extensiones Y Timezone")
    code(
        document,
        r"""
        create extension if not exists pgcrypto with schema extensions;
        create extension if not exists pg_net with schema extensions;
        create extension if not exists supabase_vault with schema vault;
        create extension if not exists pg_cron with schema pg_catalog;

        alter database postgres set timezone to 'America/Guayaquil';
        """,
        "SQL inicial en Supabase SQL Editor",
    )
    p(
        document,
        "Para proyectos nuevos, aplicar el SQL base documentado y luego las migraciones finales del repo. "
        "Si se usa docs/supabase_replicacion_simpliale_solo_public_cw.sql como base histórica, ejecutar después las migraciones 20260519130831 y 20260519133145 para dejar el modelo final sin public.dashboard_tag_settings vigente."
    )
    table(
        document,
        ["Orden", "Archivo", "Propósito"],
        [
            ["1", "docs/supabase_replicacion_simpliale_solo_public_cw.sql", "Base histórica de schema public mínimo, cw, cron, vault, roles y seeds. Reemplazar placeholders antes de ejecutar."],
            ["2", "supabase/migrations/20260519130831_auto_discover_chatwoot_catalogs.sql", "Crea catálogos automáticos de etiquetas y atributos de Chatwoot."],
            ["3", "supabase/migrations/20260519133145_centralize_dashboard_settings_in_cw.sql", "Centraliza settings en cw, hace discovery read-only y elimina public.dashboard_tag_settings."],
        ],
        [0.7, 3.3, 3.0],
    )

    h(document, 2, "9.1 Mapa De Bloques Del SQL Base")
    p(
        document,
        "El archivo docs/supabase_replicacion_simpliale_solo_public_cw.sql se ejecuta por bloques. "
        "Cada bloque tiene marcadores INICIO/FIN para copiar desde el SQL Editor. Si un bloque falla, detenerse, corregir y no avanzar."
    )
    table(
        document,
        ["Bloque", "Nombre", "Qué hace"],
        BASE_SQL_BLOCKS,
        [0.7, 2.2, 4.3],
    )

    h(document, 2, "9.2 Creación Completa De Tablas")
    p(
        document,
        "La creación completa de tablas se ejecuta desde el SQL base histórico y las migraciones finales. "
        "Para una IA implementadora, esto significa: no basta con crear solo Edge Functions; primero debe existir la estructura de base de datos."
    )
    table(
        document,
        ["Schema", "Tablas que deben quedar creadas", "Fuente"],
        [
            [
                "public",
                "user_profiles. Además se crea el tipo public.app_role con roles platform_admin, company_admin y operator. "
                "public.dashboard_tag_settings no es vigente y debe quedar eliminado por la migración final.",
                "Bloque public del SQL base + migración 20260519133145.",
            ],
            [
                "cw",
                "account_config, sync_cursor, sync_runs, raw_ingest, inboxes, teams, attribute_definitions, contacts_current, "
                "contact_inboxes, conversations_current, contact_attribute_history, conversation_attribute_history, "
                "conversation_label_history, business_stage_history, messages, reporting_events, daily_metrics, "
                "automated_reports, dashboard_tag_settings, conversation_label_events, automated_report_runs, "
                "import_batches, import_batch_errors, label_catalog, attribute_key_catalog.",
                "SQL base + migraciones finales de discovery y centralización.",
            ],
        ],
        [1.0, 4.7, 1.8],
    )
    callout(
        document,
        "Orden obligatorio",
        "Primero se crea la base de datos y el trigger de public.user_profiles. Después se crean usuarios en Authentication. "
        "El perfil no aparece en public.user_profiles hasta que el usuario existe en auth.users.",
        "yellow",
    )

    h(document, 2, "9.3 Migraciones Del Repo Y Para Qué Sirven")
    p(
        document,
        "Estas migraciones explican el historial técnico del repo. Si el SQL consolidado ya incluye parte de ellas, aun así la IA debe revisar y aplicar las migraciones finales para garantizar el estado vigente."
    )
    table(
        document,
        ["Migración", "Qué agrega/corrige", "Cuándo importa"],
        MIGRATION_DICTIONARY,
        [2.8, 2.6, 1.9],
    )

    h(document, 1, "10. Roles De Usuarios")
    p(
        document,
        "Antes de ejecutar los UPDATE de roles, crear manualmente al menos tres usuarios en Supabase Dashboard > Authentication > Users. "
        "La recomendación operativa es crear un usuario para cada rol base y luego asignar el rol en public.user_profiles."
    )
    table(
        document,
        ["Paso", "Acción", "Resultado"],
        [
            ["1", "Ir a Supabase Dashboard > Authentication > Users > Add user.", "Se crea el usuario real en auth.users."],
            ["2", "Crear mínimo tres correos: <EMAIL_PLATFORM_ADMIN>, <EMAIL_COMPANY_ADMIN>, <EMAIL_OPERATOR>.", "Cada usuario dispara el trigger hacia public.user_profiles."],
            ["3", "Confirmar que public.user_profiles tenga una fila por usuario.", "Todos nacen por defecto como operator."],
            ["4", "Actualizar roles en public.user_profiles.", "Quedan asignados platform_admin, company_admin y operator."],
            ["5", "Verificar login en dashboard con cada usuario.", "Permisos de UI, RLS y lecturas funcionan por rol."],
        ],
        [0.8, 3.0, 3.4],
    )
    table(
        document,
        ["Rol", "Ejemplo de correo", "Uso esperado"],
        [
            ["platform_admin", "admin@simplia.com", "Administrador global de la plataforma y configuración sensible."],
            ["company_admin", "test@simplia.com", "Administrador operativo de la empresa/dashboard."],
            ["operator", "operator@simplia.com", "Usuario operativo con permisos limitados."],
        ],
        [1.4, 2.0, 3.7],
    )
    code(
        document,
        r"""
        update public.user_profiles
        set role = 'platform_admin'
        where email = '<EMAIL_PLATFORM_ADMIN>';

        update public.user_profiles
        set role = 'company_admin'
        where email = '<EMAIL_COMPANY_ADMIN>';

        update public.user_profiles
        set role = 'operator'
        where email = '<EMAIL_OPERATOR>';

        select id, email, role, created_at
        from public.user_profiles
        order by role, email;
        """,
        "Asignación de roles",
    )

    h(document, 1, "11. Edge Function Secrets")
    p(document, "Total recomendado para la app: 13 secrets custom de aplicación más los secrets propios de Supabase que el runtime ya expone o que se verifican en el proyecto.")
    table(
        document,
        ["Secret", "Obligatorio", "Valor que debe ir", "Uso"],
        [
            ["VITE_CHATWOOT_BASE_URL", "Sí", "<CHATWOOT_BASE_URL>", "chatwoot-sync; base para endpoints API."],
            ["VITE_CHATWOOT_ACCOUNT_ID", "Sí", "<CHATWOOT_ACCOUNT_ID>", "chatwoot-sync; account path."],
            ["VITE_CHATWOOT_API_TOKEN", "Sí", "<CHATWOOT_API_TOKEN>", "chatwoot-sync; header api_access_token."],
            ["CHATWOOT_BASE_URL", "Sí", "Mismo que VITE_CHATWOOT_BASE_URL", "Alias backend para repair."],
            ["CHATWOOT_ACCOUNT_ID", "Sí", "Mismo que VITE_CHATWOOT_ACCOUNT_ID", "Alias backend para repair."],
            ["CHATWOOT_API_TOKEN", "Sí", "Mismo que VITE_CHATWOOT_API_TOKEN", "Alias backend para repair."],
            ["CHATWOOT_WEBHOOK_SECRET", "Sí", "Cadena aleatoria 32-64+ chars", "Valida llamadas de Chatwoot al webhook."],
            ["OPENAI_API_KEY", "Sí para IA", "<OPENAI_API_KEY>", "generate-ai-report y reportes IA programados."],
            ["OPENAI_REPORT_MODEL", "Opcional", "gpt-5.4-mini", "Modelo por defecto para reportes IA."],
            ["OPENAI_REPORT_REASONING_EFFORT", "Opcional", "low", "Esfuerzo de razonamiento para reportes IA."],
            ["RESEND_API_KEY", "Sí para correo", "<RESEND_API_KEY>", "send-scheduled-reports."],
            ["RESEND_FROM_EMAIL", "Sí para correo", "Simplia Leads <reportes@simpliaconsulting.com>", "Remitente verificado en Resend."],
            ["DASHBOARD_ACCOUNT_ID", "Opcional", "0", "Account interno de settings; default 0."],
            ["SUPABASE_URL", "Verificar", "https://<PROJECT_REF>.supabase.co", "Cliente service role dentro de funciones."],
            ["SUPABASE_SERVICE_ROLE_KEY", "Verificar", "<SUPABASE_SERVICE_ROLE_KEY>", "Escritura server-side; nunca frontend."],
            ["SUPABASE_ANON_KEY", "Verificar", "<SUPABASE_ANON_KEY>", "Invocaciones con JWT activo y frontend."],
        ],
        [1.7, 0.8, 2.4, 2.2],
    )
    code(
        document,
        r"""
        npx supabase secrets set `
          VITE_CHATWOOT_BASE_URL="<CHATWOOT_BASE_URL>" `
          VITE_CHATWOOT_ACCOUNT_ID="<CHATWOOT_ACCOUNT_ID>" `
          VITE_CHATWOOT_API_TOKEN="<CHATWOOT_API_TOKEN>" `
          CHATWOOT_BASE_URL="<CHATWOOT_BASE_URL>" `
          CHATWOOT_ACCOUNT_ID="<CHATWOOT_ACCOUNT_ID>" `
          CHATWOOT_API_TOKEN="<CHATWOOT_API_TOKEN>" `
          CHATWOOT_WEBHOOK_SECRET="<CHATWOOT_WEBHOOK_SECRET>" `
          OPENAI_API_KEY="<OPENAI_API_KEY>" `
          OPENAI_REPORT_MODEL="gpt-5.4-mini" `
          OPENAI_REPORT_REASONING_EFFORT="low" `
          RESEND_API_KEY="<RESEND_API_KEY>" `
          RESEND_FROM_EMAIL="Simplia Leads <reportes@simpliaconsulting.com>" `
          DASHBOARD_ACCOUNT_ID="0" `
          --project-ref "<PROJECT_REF>"

        npx supabase secrets list --project-ref "<PROJECT_REF>"
        """,
        "Seteo de secrets por CLI",
    )

    h(document, 1, "12. Vault Secrets Para Cron")
    p(document, "Estos secretos no son Edge Function Secrets. Viven en vault.secrets y los usa pg_cron/pg_net desde Postgres.")
    table(
        document,
        ["Vault secret", "Valor", "Uso"],
        [
            ["chatwoot_sync_project_url", "https://<PROJECT_REF>.supabase.co", "Base URL para invocar /functions/v1/chatwoot-sync y /send-scheduled-reports."],
            ["chatwoot_sync_jwt", "<SUPABASE_SERVICE_ROLE_KEY> o JWT válido", "Bearer token usado por pg_net para funciones con JWT activo."],
        ],
        [2.0, 2.4, 3.0],
    )
    code(
        document,
        r"""
        select vault.create_secret(
          'https://<PROJECT_REF>.supabase.co',
          'chatwoot_sync_project_url',
          'Supabase project URL used by pg_cron to invoke Edge Functions'
        );

        select vault.create_secret(
          '<SUPABASE_SERVICE_ROLE_KEY>',
          'chatwoot_sync_jwt',
          'JWT used by pg_cron to invoke Edge Functions'
        );
        """,
        "Ejemplo de Vault",
    )

    h(document, 1, "13. Desplegar Edge Functions")
    p(
        document,
        "Las Edge Functions son el backend server-side de la implementación. Usan service role y secrets; por eso no deben reemplazarse por llamadas directas desde el navegador cuando hay claves privadas."
    )
    table(
        document,
        ["Función", "Qué hace", "Seguridad", "Cuándo se usa"],
        EDGE_FUNCTION_DICTIONARY,
        [1.5, 2.8, 1.6, 1.6],
    )
    table(
        document,
        ["Función", "JWT", "Motivo"],
        [
            ["chatwoot-sync", "Activo", "Invocación manual/cron con JWT."],
            ["chatwoot-repair-conversations", "Activo", "Uso interno/admin con JWT."],
            ["generate-ai-report", "Activo", "Invocación desde dashboard autenticado."],
            ["send-scheduled-reports", "Activo", "Invocación por cron/Vault JWT."],
            ["chatwoot-label-webhook", "Desactivado", "Chatwoot no envía JWT Supabase; se valida con CHATWOOT_WEBHOOK_SECRET."],
        ],
        [2.2, 1.0, 4.0],
    )
    code(
        document,
        r"""
        npx supabase functions deploy chatwoot-sync chatwoot-repair-conversations generate-ai-report send-scheduled-reports --project-ref "<PROJECT_REF>" --use-api
        npx supabase functions deploy chatwoot-label-webhook --project-ref "<PROJECT_REF>" --use-api --no-verify-jwt
        npx supabase functions list --project-ref "<PROJECT_REF>"
        """,
        "Deploy",
    )

    h(document, 1, "14. Configurar Chatwoot")
    bullets(
        document,
        [
            "Validar que existan inboxes, etiquetas y custom attributes en Chatwoot. Supabase no necesita que todos los negocios tengan las mismas etiquetas.",
            "Crear webhook en Chatwoot con URL https://<PROJECT_REF>.supabase.co/functions/v1/chatwoot-label-webhook?secret=<CHATWOOT_WEBHOOK_SECRET>.",
            "Suscribir eventos: conversation_created, conversation_updated, conversation_status_changed, contact_created, contact_updated, message_created, message_updated.",
            "El valor <CHATWOOT_WEBHOOK_SECRET> debe ser idéntico al Edge Function Secret CHATWOOT_WEBHOOK_SECRET.",
            "Después de crear nuevas etiquetas o atributos, el próximo sync o webhook actualizará cw.label_catalog y cw.attribute_key_catalog.",
        ],
    )

    h(document, 1, "15. Data API")
    p(
        document,
        "El frontend usa supabase.schema(\"cw\") para leer tablas del dashboard. Si el proyecto no expone cw automáticamente, agregar cw en Supabase Dashboard > Integrations > Data API > Exposed schemas. "
        "Mantener RLS y grants como defensa; exponer un schema no sustituye políticas de seguridad."
    )
    table(
        document,
        ["Validación", "Esperado"],
        [
            ["cw expuesto en Data API", "El dashboard puede leer cw.conversations_current, cw.dashboard_tag_settings, cw.label_catalog y cw.attribute_definitions."],
            ["RLS habilitado", "Las tablas de cw y public.user_profiles tienen políticas adecuadas."],
            ["No fallback public settings", "El código lee cw.dashboard_tag_settings; public.dashboard_tag_settings no existe o no se usa."],
        ],
        [2.4, 4.8],
    )

    h(document, 1, "16. Cron Operativo")
    table(
        document,
        ["Job", "Schedule", "Equivalencia", "Función"],
        [
            ["sync-chatwoot-diario", "1 5 * * *", "00:01 America/Guayaquil porque pg_cron usa GMT", "chatwoot-sync"],
            ["send-scheduled-reports", "*/5 * * * *", "Cada 5 minutos", "send-scheduled-reports"],
        ],
        [1.8, 1.2, 2.4, 1.6],
    )
    code(
        document,
        r"""
        select current_setting('TIMEZONE') as db_timezone,
               current_setting('cron.timezone', true) as cron_timezone;

        select jobid, jobname, schedule, active
        from cron.job
        where jobname in ('sync-chatwoot-diario', 'send-scheduled-reports');

        select d.jobid, j.jobname, d.runid, d.status, d.return_message, d.start_time, d.end_time
        from cron.job_run_details d
        join cron.job j on j.jobid = d.jobid
        where j.jobname in ('sync-chatwoot-diario', 'send-scheduled-reports')
        order by d.start_time desc
        limit 20;
        """,
        "Validar cron",
    )

    h(document, 1, "17. Pruebas De Implementación")
    table(
        document,
        ["Prueba", "Comando / acción", "Éxito esperado"],
        [
            ["Secrets", "npx supabase secrets list --project-ref <PROJECT_REF>", "Aparecen los nombres esperados. No se ven valores planos."],
            ["Funciones", "npx supabase functions list --project-ref <PROJECT_REF>", "Todas ACTIVE; webhook desplegado sin JWT."],
            ["Sync inicial", "POST /functions/v1/chatwoot-sync con window_hours 72", "cw.sync_runs status success; stats con labels, attribute_definitions, contacts y conversations."],
            ["Catálogos", "select count(*) from cw.label_catalog / cw.attribute_definitions / cw.attribute_key_catalog", "Conteos mayores a 0 si Chatwoot tiene datos."],
            ["Settings", "select to_regclass('public.dashboard_tag_settings'), to_regclass('cw.dashboard_tag_settings')", "public null; cw existe."],
            ["Reportes IA", "Invocar generate-ai-report desde dashboard", "Respuesta 200 o job iniciado; logs sin Missing OPENAI_API_KEY."],
            ["Reportes correo", "Invocar send-scheduled-reports", "Respuesta 200; Resend no rechaza remitente."],
            ["Webhook", "Cambiar etiqueta/contacto en Chatwoot", "cw.conversation_label_events o snapshots se actualizan; logs 200."],
            ["Cron", "Revisar cron.job_run_details y cw.sync_health", "Jobs activos y sin errores pg_net."],
        ],
        [1.5, 3.1, 2.6],
    )
    code(
        document,
        r"""
        $body = @{
          mode = 'full'
          window_hours = 72
          sync_messages = 'recent'
          dashboard_account_id = 0
        } | ConvertTo-Json -Compress

        Invoke-RestMethod `
          -Method Post `
          -Uri "$env:VITE_SUPABASE_URL/functions/v1/chatwoot-sync" `
          -Headers @{ Authorization = "Bearer $env:VITE_SUPABASE_ANON_KEY"; apikey = $env:VITE_SUPABASE_ANON_KEY } `
          -ContentType 'application/json' `
          -Body $body `
          -TimeoutSec 120
        """,
        "Prueba manual de sync",
    )

    h(document, 1, "18. Limpieza De Datos De Prueba")
    bullets(
        document,
        [
            "Crear datos de prueba con prefijo claro, por ejemplo TEST_IA_ o test@example.com.",
            "Eliminar primero la conversación/contacto de prueba en Chatwoot si se creó solo para validación.",
            "No borrar datos reales desde Supabase. Si se limpia Supabase, hacerlo solo por chatwoot_conversation_id o chatwoot_contact_id de prueba confirmado.",
            "Después de limpiar, ejecutar chatwoot-sync o chatwoot-repair-conversations para reconciliar snapshots.",
        ],
    )

    h(document, 1, "19. Troubleshooting")
    table(
        document,
        ["Síntoma", "Causa probable", "Qué revisar"],
        [
            ["Missing required Supabase or Chatwoot environment variables", "Faltan secrets VITE_CHATWOOT_* o SUPABASE_*.", "npx supabase secrets list y nombres exactos."],
            ["Webhook 401", "CHATWOOT_WEBHOOK_SECRET no coincide.", "URL de webhook en Chatwoot y secret en Supabase."],
            ["Cron activo pero no hay sync_runs", "Aún no llegó la hora o pg_net/JWT falló.", "cron.job_run_details, net._http_response, cw.sync_health."],
            ["Resend rechaza correo", "Dominio/remitente no verificado.", "RESEND_FROM_EMAIL y panel Resend."],
            ["OpenAI falla", "API key ausente, modelo inválido o timeout.", "OPENAI_API_KEY, OPENAI_REPORT_MODEL y logs de Edge Function."],
            ["Dashboard no lee cw", "Schema cw no está expuesto o faltan grants/RLS.", "Data API settings, policies y consola del navegador."],
            ["Etiquetas nuevas no salen", "Aún no corrió sync/webhook o Chatwoot no las tiene.", "chatwoot-sync, label_catalog, webhook logs."],
        ],
        [1.8, 2.3, 3.0],
    )

    h(document, 1, "20. Checklist Para Una IA Implementadora")
    numbered(
        document,
        [
            "Leer este documento completo y listar valores faltantes sin pedir secretos que ya estén en .env o Supabase CLI.",
            "Verificar .env y Supabase CLI antes de aplicar cambios.",
            "No imprimir valores secretos en logs o respuesta final.",
            "Ejecutar SQL por bloques y detenerse ante el primer error.",
            "Crear/validar secrets con placeholders reales entregados por el usuario.",
            "Desplegar funciones con el flag JWT correcto.",
            "Configurar Chatwoot webhook con secret compartido.",
            "Ejecutar sync inicial y pruebas de catálogos.",
            "Verificar cron y logs.",
            "Entregar evidencia de comandos exitosos, tablas validadas y riesgos restantes.",
        ],
    )

    h(document, 1, "21. Anexo De Comandos Rápidos")
    code(
        document,
        r"""
        # CLI
        npx supabase --version
        npx supabase login --token "<SUPABASE_ACCESS_TOKEN>"
        npx supabase link --project-ref "<PROJECT_REF>"
        npx supabase secrets list --project-ref "<PROJECT_REF>"

        # Deploy
        npx supabase functions deploy chatwoot-sync chatwoot-repair-conversations generate-ai-report send-scheduled-reports --project-ref "<PROJECT_REF>" --use-api
        npx supabase functions deploy chatwoot-label-webhook --project-ref "<PROJECT_REF>" --use-api --no-verify-jwt

        # Validación SQL
        select to_regclass('public.dashboard_tag_settings') as public_table,
               to_regclass('cw.dashboard_tag_settings') as cw_table;

        select count(*) from cw.label_catalog;
        select count(*) from cw.attribute_definitions;
        select count(*) from cw.attribute_key_catalog;
        select * from cw.sync_health;
        """,
    )

    h(document, 1, "22. Anexo SQL Actualizado")
    p(
        document,
        "El SQL base histórico sigue disponible en el repo, pero la fuente de verdad para el estado final incluye estas migraciones posteriores. "
        "No copiar el SQL viejo como verdad final sin aplicar estos parches."
    )
    table(
        document,
        ["Archivo", "Estado"],
        [
            [str(LEGACY_SQL.relative_to(REPO_ROOT)), "Base histórica. Útil como arranque, pero contiene compatibilidad antigua que debe corregirse."],
            ["supabase/migrations/20260519130831_auto_discover_chatwoot_catalogs.sql", "Crear/actualizar discovery automático de etiquetas y atributos."],
            ["supabase/migrations/20260519133145_centralize_dashboard_settings_in_cw.sql", "Centralizar settings en cw y eliminar tabla public legacy."],
        ],
        [3.6, 3.6],
    )
    callout(
        document,
        "Cómo usar este anexo",
        "Primero copiar/ejecutar el SQL base por bloques. Después ejecutar la migración de discovery automático y finalmente la migración de centralización en cw. "
        "Así se conserva todo lo necesario de cw y se elimina la compatibilidad vieja de public.dashboard_tag_settings.",
        "yellow",
    )
    code(
        document,
        read_file(LEGACY_SQL),
        "SQL base consolidado: crear public.user_profiles, schema cw, RLS, triggers, Vault, cron, roles y validaciones",
    )
    document.add_section(WD_SECTION.NEW_PAGE)
    code(
        document,
        read_file(MIGRATIONS_DIR / "20260519130831_auto_discover_chatwoot_catalogs.sql"),
        "Migración final 1: discovery automático",
    )
    document.add_section(WD_SECTION.NEW_PAGE)
    code(
        document,
        read_file(MIGRATIONS_DIR / "20260519133145_centralize_dashboard_settings_in_cw.sql"),
        "Migración final 2: centralización en cw",
    )

    document.save(OUTPUT_DOCX)


if __name__ == "__main__":
    build_document()
    print(f"Generated {OUTPUT_DOCX}")
